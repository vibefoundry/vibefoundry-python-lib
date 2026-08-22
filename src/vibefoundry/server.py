"""
FastAPI backend server for VibeFoundry IDE
"""

import os
import sys
import json
import math
import re
import asyncio
import base64
import secrets
import struct
import shutil
import signal
import time
import urllib.parse
import webbrowser
from datetime import datetime, timezone
from pathlib import Path

# Safe despite __init__ importing cli: __version__ is bound before that import,
# so a partially-initialized package still resolves it.
from vibefoundry import __version__
from vibefoundry import xlsx_view

# Honor the OS-native trust store (Windows cert store, macOS Keychain) so
# corporate TLS-inspecting proxies — which re-sign traffic with an internal
# CA that lives only in the OS store, not certifi — don't break HTTPS calls
# to vibefoundry.ai. Must run before any SSLContext / httpx client is built.
try:
    import truststore
    truststore.inject_into_ssl()
except ImportError:
    pass


def _safe_float_or_none(v):
    """Coerce to JSON-safe float or None. NaN/Inf/None all become None (shown blank in UI)."""
    if v is None:
        return None
    try:
        f = float(v)
    except (TypeError, ValueError):
        return None
    if math.isnan(f) or math.isinf(f):
        return None
    return f


def _safe_int(v):
    """Coerce to int. None/NaN become 0. Use for count-style stats that are always integers ≥ 0."""
    if v is None:
        return 0
    try:
        f = float(v)
    except (TypeError, ValueError):
        return 0
    if math.isnan(f) or math.isinf(f):
        return 0
    return int(f)

# Unix-only imports for terminal functionality
if sys.platform != 'win32':
    import pty
    import fcntl
    import termios
    import select
else:
    pty = None
    fcntl = None
    termios = None
    select = None
from typing import Optional
from contextlib import asynccontextmanager

import httpx
import polars as pl
from fastapi import FastAPI, HTTPException, Request, WebSocket, WebSocketDisconnect, UploadFile, File, Form, Query
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse, Response
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

from vibefoundry.runner import discover_scripts, run_script, setup_project_structure, ScriptResult, stop_all_scripts, list_running_processes, stop_process
from vibefoundry.metadata import generate_metadata
from vibefoundry.watcher import FileWatcher
from vibefoundry.organizations import (
    ORGANIZATIONS, PUBLIC_ORG_ID, PUBLIC_ORG_NAME,
    find_organization, normalize_hub_url, normalize_gateway_url, org_hint_from_hub_url,
)
from vibefoundry.profiler import (
    is_file_massive, get_profile_cache_path, is_profile_valid,
    profile_large_file, read_cached_profile, estimate_filtered_rows,
    apply_column_exclusions,
    _detect_csv_separator, _get_lazy_frame,
)


# Global state
class AppState:
    project_folder: Optional[Path] = None
    watcher: Optional[FileWatcher] = None
    websocket_clients: list[WebSocket] = []
    # Debounce for script change notifications (prevent duplicates)
    last_script_change: dict[str, float] = {}  # path -> timestamp
    # True when a host plugin registered this backend as an embedded pane.
    # Set via POST /api/ui/pane by code that KNOWS (the plugin, at open time),
    # read by the frontend at boot — so pane chrome is a fact served by the
    # backend, not a guess from URLs or embedding tricks. Claude's preview is
    # a native webview (not framed) and forbids query strings in its config
    # URLs, which killed every client-side detection in turn.
    pane_mode: bool = False
    # True when something outside the UI wants the Organization Catalogue on
    # screen — the plugin sets it from connect_organization, because a tool
    # call has no way to reach into the already-loaded frontend. Read once at
    # boot and by the health poll, same contract as pane_mode.
    open_org_catalog: bool = False


class DataFrameState:
    """Stream-from-disk DataFrame viewer - only loads rows as needed"""
    def __init__(self):
        self.file_path: Optional[str] = None
        self.file_type: Optional[str] = None  # 'csv' or 'excel'
        self.csv_separator: str = ','
        self.columns: list[str] = []
        self.column_info: dict = {}  # {col: {type, min, max, values}}
        self.total_rows: int = 0
        self.current_filters: dict = {}
        self.current_sort: Optional[dict] = None
        self.row_limit: Optional[int] = None  # Cap rows for large file preview
        # Small cache for filtered row count (avoids re-scanning)
        self._filtered_row_count: Optional[int] = None

    def clear(self):
        """Clear state"""
        print(f"[Memory] Clearing DataFrame state")
        self.file_path = None
        self.file_type = None
        self.csv_separator = ','
        self.columns = []
        self.column_info = {}
        self.total_rows = 0
        self.current_filters = {}
        self.current_sort = None
        self.row_limit = None
        self._filtered_row_count = None

    def _get_lazy_frame(self) -> Optional[pl.LazyFrame]:
        """Get a lazy frame for the file (doesn't load data)"""
        if not self.file_path:
            return None
        file_path = Path(self.file_path)
        if self.file_type == 'csv':
            return pl.scan_csv(file_path, separator=self.csv_separator, infer_schema_length=10000)
        elif self.file_type == 'parquet':
            return pl.scan_parquet(file_path)
        elif self.file_type == 'excel':
            # Excel doesn't support lazy scanning, load eagerly but this is rare
            return pl.read_excel(file_path).lazy()
        return None

    def _apply_filters_sort(self, lf: pl.LazyFrame) -> pl.LazyFrame:
        """Apply current filters and sort to a lazy frame"""
        try:
            schema = lf.collect_schema()
        except Exception:
            schema = None

        # Apply filters
        for column, filter_val in self.current_filters.items():
            if column not in self.columns:
                continue
            if isinstance(filter_val, dict):
                if 'values' in filter_val:
                    # Categorical filter (object form with optional exclude)
                    vals = filter_val.get('values') or []
                    if vals:
                        str_vals = [str(v) for v in vals]
                        lf = lf.filter(pl.col(column).cast(pl.Utf8).is_in(str_vals))
                else:
                    # Numeric range filter
                    if filter_val.get('min') not in (None, '', 'null'):
                        try:
                            min_val = float(filter_val['min'])
                            lf = lf.filter(pl.col(column).cast(pl.Float64, strict=False) >= min_val)
                        except (ValueError, TypeError):
                            pass
                    if filter_val.get('max') not in (None, '', 'null'):
                        try:
                            max_val = float(filter_val['max'])
                            lf = lf.filter(pl.col(column).cast(pl.Float64, strict=False) <= max_val)
                        except (ValueError, TypeError):
                            pass
                lf = apply_column_exclusions(lf, column, filter_val.get('exclude') or [], schema)
            elif isinstance(filter_val, list) and len(filter_val) > 0:
                # Categorical filter with special sentinels for null/blank/zero
                SPECIAL = {'__vf_filter_null__', '__vf_filter_blank__', '__vf_filter_zero__'}
                specials = [v for v in filter_val if v in SPECIAL]
                regular = [v for v in filter_val if v not in SPECIAL]
                predicates = []
                if regular:
                    str_vals = [str(v) for v in regular]
                    predicates.append(pl.col(column).cast(pl.Utf8).is_in(str_vals))
                if '__vf_filter_null__' in specials:
                    predicates.append(pl.col(column).is_null())
                if '__vf_filter_blank__' in specials:
                    predicates.append(pl.col(column).cast(pl.Utf8) == '')
                if '__vf_filter_zero__' in specials:
                    predicates.append(pl.col(column).cast(pl.Float64, strict=False) == 0)
                if predicates:
                    combined = predicates[0]
                    for p in predicates[1:]:
                        combined = combined | p
                    lf = lf.filter(combined)

        # Apply sort
        if self.current_sort and self.current_sort.get('column'):
            sort_col = self.current_sort['column']
            descending = self.current_sort.get('direction', 'asc') != 'asc'
            if sort_col in self.columns:
                lf = lf.sort(sort_col, descending=descending, nulls_last=True)

        return lf

    def get_rows(self, offset: int, limit: int) -> tuple[list[dict], int]:
        """Get rows with current filters/sort applied. Returns (rows, total_filtered_count)"""
        lf = self._get_lazy_frame()
        if lf is None:
            return [], 0

        lf = self._apply_filters_sort(lf)

        # Apply row limit if set (large file preview)
        if self.row_limit is not None:
            lf = lf.slice(0, self.row_limit)

        # Get total count (cached if no filter changes)
        if self._filtered_row_count is None:
            self._filtered_row_count = lf.select(pl.len()).collect().item()

        # Get requested slice
        rows_df = lf.slice(offset, limit).collect()
        rows = rows_df.to_dicts()

        # Null → blank, NaN → "NaN" (visible), Inf → blank (for JSON safety)
        for row in rows:
            for key in row:
                v = row[key]
                if v is None:
                    row[key] = ''
                elif isinstance(v, float):
                    if math.isnan(v):
                        row[key] = 'NaN'
                    elif math.isinf(v):
                        row[key] = ''

        return rows, self._filtered_row_count

    def invalidate_filter_cache(self):
        """Call when filters change"""
        self._filtered_row_count = None


state = AppState()
df_state = DataFrameState()

# Track active profiling task so it can be cancelled
_profiling_task: Optional[asyncio.Task] = None


def _compute_column_info(lf: pl.LazyFrame, columns: list, schema) -> dict:
    """Compute column info in a single optimized pass.
    Batches all numeric stats into one query, then handles categorical columns.
    Returns min/max/nullCount/zeroCount for numeric, values/nullCount/blankCount for categorical."""
    column_info = {}

    # Separate numeric and categorical columns
    numeric_cols = []
    categorical_cols = []
    for col in columns:
        dtype = schema.get(col)
        if dtype is None:
            continue
        if dtype.is_numeric():
            numeric_cols.append(col)
        else:
            categorical_cols.append(col)

    # Batch all numeric column stats in ONE query (single file scan)
    if numeric_cols:
        try:
            float_cols = {c for c in numeric_cols if schema.get(c) in (pl.Float32, pl.Float64)}
            exprs = []
            for col in numeric_cols:
                # fill_nan(None) so NaN is ignored in aggregations (treated as missing)
                clean = pl.col(col).fill_nan(None) if col in float_cols else pl.col(col)
                exprs.extend([
                    clean.min().alias(f'{col}__min'),
                    clean.max().alias(f'{col}__max'),
                    clean.sum().alias(f'{col}__sum'),
                    clean.mean().alias(f'{col}__mean'),
                    clean.median().alias(f'{col}__median'),
                    pl.col(col).count().alias(f'{col}__count'),
                    pl.col(col).is_null().sum().alias(f'{col}__null'),
                    (pl.col(col) == 0).sum().alias(f'{col}__zero'),
                    pl.col(col).drop_nulls().n_unique().alias(f'{col}__unique'),
                ])
                if col in float_cols:
                    exprs.append(pl.col(col).is_nan().sum().alias(f'{col}__nan'))
            stats = lf.select(exprs).collect()

            for col in numeric_cols:
                column_info[col] = {
                    "type": "numeric",
                    "min": _safe_float_or_none(stats[f'{col}__min'][0]),
                    "max": _safe_float_or_none(stats[f'{col}__max'][0]),
                    "sum": _safe_float_or_none(stats[f'{col}__sum'][0]),
                    "mean": _safe_float_or_none(stats[f'{col}__mean'][0]),
                    "median": _safe_float_or_none(stats[f'{col}__median'][0]),
                    "count": _safe_int(stats[f'{col}__count'][0]),
                    "nullCount": _safe_int(stats[f'{col}__null'][0]),
                    "zeroCount": _safe_int(stats[f'{col}__zero'][0]),
                    "uniqueCount": _safe_int(stats[f'{col}__unique'][0]),
                    "nanCount": _safe_int(stats[f'{col}__nan'][0]) if col in float_cols else 0,
                }
        except Exception:
            for col in numeric_cols:
                column_info[col] = {"type": "numeric", "min": None, "max": None, "sum": None, "mean": None, "median": None, "count": 0, "nullCount": 0, "zeroCount": 0, "uniqueCount": 0, "nanCount": 0}

    # Batch categorical stats in ONE query
    if categorical_cols:
        try:
            exprs = []
            for col in categorical_cols:
                exprs.extend([
                    pl.col(col).count().alias(f'{col}__count'),
                    pl.col(col).is_null().sum().alias(f'{col}__null'),
                    (pl.col(col).cast(pl.Utf8) == '').sum().alias(f'{col}__blank'),
                    pl.col(col).drop_nulls().n_unique().alias(f'{col}__unique'),
                ])
            stats = lf.select(exprs).collect()

            # Get unique values for each categorical column (requires separate queries for .unique())
            for col in categorical_cols:
                try:
                    unique_vals = lf.select(
                        pl.col(col).drop_nulls().cast(pl.Utf8).unique()
                    ).collect()[col].to_list()
                    unique_vals = sorted([str(v) for v in unique_vals if v != ''])
                except Exception:
                    unique_vals = []

                column_info[col] = {
                    "type": "categorical",
                    "values": unique_vals,
                    "count": _safe_int(stats[f'{col}__count'][0]),
                    "nullCount": _safe_int(stats[f'{col}__null'][0]),
                    "blankCount": _safe_int(stats[f'{col}__blank'][0]),
                    "uniqueCount": _safe_int(stats[f'{col}__unique'][0]),
                }
        except Exception:
            for col in categorical_cols:
                column_info[col] = {"type": "categorical", "values": [], "count": 0, "nullCount": 0, "blankCount": 0, "uniqueCount": 0}

    return column_info


# Alias for backward compatibility
_compute_full_column_info = _compute_column_info


# Request/Response models
class FolderSelectRequest(BaseModel):
    path: str


class RunScriptsRequest(BaseModel):
    scripts: list[str]


class ScriptResultResponse(BaseModel):
    script_path: str
    success: bool
    stdout: str
    stderr: str
    return_code: int
    error: Optional[str] = None
    timed_out: bool = False
    streamlit_url: Optional[str] = None  # URL if this was a Streamlit app


@asynccontextmanager
async def lifespan(app: FastAPI):
    """Startup and shutdown events"""
    # Check for project folder from environment
    project_path = os.environ.get("VIBEFOUNDRY_PROJECT_PATH")
    if project_path:
        folder = Path(project_path)
        if folder.exists() and folder.is_dir():
            state.project_folder = folder
            # No scaffolding here: opening a folder must not modify it. Every
            # other path already lives by that rule — generate_metadata no-ops
            # until app_folder exists, the watcher only watches folders that
            # exist, folder/select creates nothing. Structure is created in
            # exactly one place: the Build button (/api/build).
            generate_metadata(folder)
            state.watcher = FileWatcher(folder)
            state.watcher.scan_initial_state()

    yield
    # Cleanup
    if state.watcher:
        state.watcher.stop()
    # Stop any running scripts (including Streamlit apps)
    stopped = stop_all_scripts()
    if stopped:
        print(f"[Shutdown] Stopped {stopped} running script(s)")


# Create FastAPI app
app = FastAPI(
    title="VibeFoundry IDE",
    version=__version__,
    lifespan=lifespan
)

# CORS for development
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


def get_static_dir() -> Path:
    """Get the path to bundled static files"""
    return Path(__file__).parent / "static"


# API Routes

@app.get("/api/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "ok",
        "version": __version__,
        "project_folder": str(state.project_folder) if state.project_folder else None,
        "pane_mode": state.pane_mode,
        "open_org_catalog": state.open_org_catalog,
    }


class PaneModeRequest(BaseModel):
    enabled: bool = True


@app.post("/api/ui/pane")
async def set_pane_mode(request: PaneModeRequest):
    """Mark this backend as pane-hosted (called by the host plugin, not the UI)."""
    state.pane_mode = bool(request.enabled)
    return {"status": "ok", "pane_mode": state.pane_mode}


@app.post("/api/ui/org-catalog")
async def set_open_org_catalog(request: PaneModeRequest):
    """Ask the UI to open the Organization Catalogue (called by the plugin)."""
    state.open_org_catalog = bool(request.enabled)
    return {"status": "ok", "open_org_catalog": state.open_org_catalog}


class LaunchTerminalRequest(BaseModel):
    path: str
    command: str = None  # Optional command to run after cd (e.g., 'claude', 'codex')


@app.post("/api/terminal/launch")
async def launch_native_terminal(request: LaunchTerminalRequest):
    """Launch a native terminal window, cd into the project, and optionally run a command"""
    import subprocess

    folder_path = Path(request.path)
    if not folder_path.exists():
        raise HTTPException(status_code=400, detail="Folder does not exist")

    if sys.platform == 'darwin':  # macOS
        # Build the bash command, then escape it once for embedding in an
        # AppleScript string literal. request.command can contain double
        # quotes — e.g. python "/path with spaces/app.py" — which would
        # otherwise terminate the AppleScript string early and break the run.
        bash_cmd = f'cd "{folder_path}" && clear'
        if request.command:
            bash_cmd += f' && {request.command}'
        escaped = bash_cmd.replace('\\', '\\\\').replace('"', '\\"')
        script = (
            'tell application "Terminal"\n'
            '    activate\n'
            f'    do script "{escaped}"\n'
            'end tell'
        )
        subprocess.run(['osascript', '-e', script], check=True)
        return {"status": "ok", "message": "Terminal launched"}
    elif sys.platform == 'win32':  # Windows
        # Use start command with /d to set working directory
        if request.command:
            subprocess.Popen(
                f'start "" /d "{folder_path}" cmd /k {request.command}',
                shell=True
            )
        else:
            subprocess.Popen(
                f'start "" /d "{folder_path}" cmd',
                shell=True
            )
        return {"status": "ok", "message": "Terminal launched"}
    else:
        raise HTTPException(status_code=400, detail="Native terminal launch not supported on this platform")


@app.post("/api/folder/select")
async def select_folder(request: FolderSelectRequest):
    """Set the project folder and initialize structure"""
    folder_path = Path(request.path)

    if not folder_path.exists():
        raise HTTPException(status_code=400, detail="Folder does not exist")

    if not folder_path.is_dir():
        raise HTTPException(status_code=400, detail="Path is not a directory")

    state.project_folder = folder_path

    # Don't auto-scaffold - user must click Build button
    # Just ensure basic folders exist for watcher
    folders = {
        "input_folder": folder_path / "input_folder",
        "output_folder": folder_path / "output_folder",
        "app_folder": folder_path / "app_folder",
        "scripts_folder": folder_path / "app_folder" / "scripts",
    }

    # Stop existing watcher
    if state.watcher:
        state.watcher.stop()

    # Start new watcher
    # Note: Pass coroutines directly - watcher.py handles thread-safe scheduling
    state.watcher = FileWatcher(
        folder_path,
        on_data_change=notify_data_change,
        on_script_change=notify_script_change,
        on_output_file_change=notify_output_file_change
    )
    await state.watcher.start_async()

    # Generate initial metadata
    generate_metadata(folder_path)

    return {
        "success": True,
        "name": folder_path.name,
        "project_folder": str(folder_path),
        "folders": {k: str(v) for k, v in folders.items()}
    }


@app.get("/api/folder/info")
async def get_folder_info():
    """Get current project folder info"""
    if not state.project_folder:
        return {"project_folder": None}

    return {
        "project_folder": str(state.project_folder),
        "name": state.project_folder.name
    }


PROXY_BASE_URL = "https://vibefoundry.ai/api/templates"
PUBLIC_TEMPLATE_FALLBACK_URL = "https://vibefoundry.ai/templates"
AUTH_LANDING_URL = "https://vibefoundry.ai/ide-auth"


# --- IDE auth (delegated Clerk Production via vibefoundry.ai) -----------------
#
# The IDE runs on http://127.0.0.1:<port>, which Clerk Production refuses to
# talk to. So sign-in is delegated: the IDE opens a browser to
#   https://vibefoundry.ai/ide-auth?state=<csrf>&callback=http://127.0.0.1:<port>/auth/callback
# the user signs in via Clerk Production there, the website mints a custom
# HMAC JWT and redirects back to the localhost callback below, which stores
# the token. Subsequent /api/build calls read the token off disk and pass
# it to the templates proxy on vibefoundry.ai.

# In-memory CSRF store: maps state nonce → unix-ts expiry. Cleaned lazily.
_pending_auth_states: dict[str, float] = {}
_AUTH_STATE_TTL_SECONDS = 600  # 10 minutes


def _auth_token_path() -> Path:
    home = Path.home() / ".vibefoundry"
    home.mkdir(parents=True, exist_ok=True)
    return home / "auth.json"


def _read_stored_token() -> Optional[dict]:
    """Returns {"token": str, "sub": str, "email": str|None, "expiresAt": int}
    or None if no stored token / expired / unreadable."""
    path = _auth_token_path()
    if not path.exists():
        return None
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return None
    expires_at = data.get("expiresAt", 0)
    if not isinstance(expires_at, (int, float)) or expires_at < time.time():
        return None
    if not data.get("token"):
        return None
    return data


def _decode_jwt_payload(token: str) -> dict:
    """Decode a JWT payload WITHOUT verifying the signature. The IDE
    doesn't have IDE_AUTH_SECRET (only the proxy does), so we trust the
    token blindly here. The proxy validates on every request."""
    try:
        parts = token.split(".")
        if len(parts) != 3:
            return {}
        payload_b64 = parts[1] + "=" * (-len(parts[1]) % 4)
        payload_bytes = base64.urlsafe_b64decode(payload_b64)
        return json.loads(payload_bytes.decode("utf-8"))
    except Exception:
        return {}


def _purge_expired_auth_states() -> None:
    now = time.time()
    expired = [s for s, exp in _pending_auth_states.items() if exp < now]
    for s in expired:
        _pending_auth_states.pop(s, None)


@app.post("/api/auth/start")
async def auth_start(request: Request):
    """Generate a CSRF state and the URL the IDE should open in the browser."""
    _purge_expired_auth_states()
    state_nonce = secrets.token_urlsafe(32)
    _pending_auth_states[state_nonce] = time.time() + _AUTH_STATE_TTL_SECONDS

    host = request.headers.get("host", "127.0.0.1:8765")
    callback_url = f"http://{host}/auth/callback"
    encoded_callback = urllib.parse.quote(callback_url, safe="")

    auth_url = (
        f"{AUTH_LANDING_URL}"
        f"?state={state_nonce}"
        f"&callback={encoded_callback}"
    )
    return {"url": auth_url}


@app.get("/auth/callback")
async def auth_callback(token: str = "", state: str = ""):
    """Receives the redirect from vibefoundry.ai/ide-auth with token+state.
    Validates state (CSRF), stores token to disk, returns an HTML page."""
    if not token or not state:
        return HTMLResponse(_auth_callback_html(error="Missing token or state in callback URL."), status_code=400)

    _purge_expired_auth_states()
    if state not in _pending_auth_states:
        return HTMLResponse(_auth_callback_html(error="Invalid or expired sign-in attempt. Try again from the IDE."), status_code=400)
    _pending_auth_states.pop(state, None)

    payload = _decode_jwt_payload(token)
    expires_at = int(payload.get("exp", 0))
    if expires_at < time.time():
        return HTMLResponse(_auth_callback_html(error="Token already expired."), status_code=400)

    record = {
        "token": token,
        "sub": payload.get("sub"),
        "email": payload.get("email"),
        "expiresAt": expires_at,
    }
    _auth_token_path().write_text(json.dumps(record), encoding="utf-8")
    return HTMLResponse(_auth_callback_html())


@app.get("/api/auth/status")
async def auth_status():
    """Read-only — frontend polls this to know if user is signed in."""
    record = _read_stored_token()
    if not record:
        return {"signedIn": False}
    return {
        "signedIn": True,
        "user": {"sub": record.get("sub"), "email": record.get("email")},
        "expiresAt": record.get("expiresAt"),
    }


@app.post("/api/auth/sign-out")
async def auth_sign_out():
    """Delete the stored token."""
    path = _auth_token_path()
    if path.exists():
        path.unlink()
    return {"signedIn": False}


def _auth_callback_html(
    error: str = "",
    page_title: str = "VibeFoundry — Sign-in complete",
    heading: str = "Signed in!",
    failed_heading: str = "Sign-in failed",
) -> str:
    # The headings are parameters so /org/callback can land on the same page
    # with its own wording instead of telling a user who just connected an
    # organization that they "signed in".
    if error:
        body = f'<p class="msg" style="color:#dc2626">{error}</p>'
    else:
        body = (
            '<p class="msg">You can close this tab and return to the VibeFoundry IDE.</p>'
            '<script>setTimeout(() => { try { window.close() } catch (e) {} }, 1500)</script>'
        )
    return f"""<!doctype html>
<html><head><title>{page_title}</title>
<style>
  body {{ font-family: system-ui, sans-serif; background: #ffffff;
         background-image: linear-gradient(rgba(37,99,235,0.11) 1px, transparent 1px),
                           linear-gradient(90deg, rgba(37,99,235,0.11) 1px, transparent 1px);
         background-size: 32px 32px;
         display: flex; align-items: center; justify-content: center;
         min-height: 100vh; margin: 0; }}
  .card {{ background: rgba(219,234,254,0.30); padding: 40px 56px;
           border: 1px solid rgba(147,197,253,0.5); border-radius: 12px;
           text-align: center; max-width: 460px; }}
  h1 {{ font-size: 22px; color: #0f172a; margin: 0 0 12px; }}
  .msg {{ font-size: 14px; color: #475569; margin: 0; }}
</style></head>
<body><div class="card"><h1>{failed_heading if error else heading}</h1>{body}</div></body></html>
"""


# --- Templates proxy + Build endpoint ----------------------------------------


async def _cascade_templates_via_proxy(dest_root: Path, jwt: str, subpath: str = "") -> list[str]:
    """Recursively fetch templates/ (including subfolders) from the private
    proxy using a Clerk JWT and write everything into dest_root. Subfolder
    structure is preserved. Raises on proxy errors so the caller can decide
    whether to fall back.

    If subpath is given (e.g. "dashboard_pwa_duckdb"), only that subtree is fetched,
    and files land at dest_root/<subpath>/... — the subpath name is preserved.
    """
    auth_headers = {"Authorization": f"Bearer {jwt}"}
    # vnd.github.raw is required so binary files (.wasm, .parquet) come back
    # as raw bytes rather than the default base64-wrapped Contents API JSON.
    file_headers = {**auth_headers, "Accept": "application/vnd.github.raw"}
    written: list[str] = []
    failed: list[tuple[str, str]] = []
    timeout = httpx.Timeout(connect=10.0, read=120.0, write=30.0, pool=10.0)

    async with httpx.AsyncClient(timeout=timeout) as client:
        async def fetch_dir(rel_path: str = "") -> None:
            url = f"{PROXY_BASE_URL}/{rel_path}" if rel_path else PROXY_BASE_URL
            listing_res = await client.get(url, headers=auth_headers)
            listing_res.raise_for_status()
            entries = listing_res.json()
            for entry in entries:
                name = entry["name"]
                entry_rel = f"{rel_path}/{name}" if rel_path else name
                if entry.get("type") == "dir":
                    (dest_root / entry_rel).mkdir(parents=True, exist_ok=True)
                    try:
                        await fetch_dir(entry_rel)
                    except Exception as e:
                        # One subtree failing shouldn't stop the rest of the cascade.
                        failed.append((entry_rel, repr(e)))
                        print(f"[Build] cascade: subtree {entry_rel} failed: {e}")
                elif entry.get("type") == "file":
                    try:
                        file_res = await client.get(f"{PROXY_BASE_URL}/{entry_rel}", headers=file_headers)
                        file_res.raise_for_status()
                        local_path = dest_root / entry_rel
                        local_path.parent.mkdir(parents=True, exist_ok=True)
                        local_path.write_bytes(file_res.content)
                        written.append(entry_rel)
                    except Exception as e:
                        # Skip-and-log so a single bad file (timeout on a big
                        # binary, transient proxy 5xx, etc.) doesn't strand
                        # the rest of the cascade — including the small
                        # files that come after it in iteration order.
                        failed.append((entry_rel, repr(e)))
                        print(f"[Build] cascade: file {entry_rel} failed: {e}")

        if subpath:
            (dest_root / subpath).mkdir(parents=True, exist_ok=True)
        await fetch_dir(subpath)
    if failed:
        print(f"[Build] cascade: {len(failed)} entries failed; {len(written)} succeeded")
    return written


# Written to the project root when Build runs and no `.env` is there yet. The
# remote copy at IDE-Agents/env.template wins when it's reachable, so the
# gateway URL and the key list can change without a package release.
#
# No blank VF_APP_ID / VF_APP_KEY here on purpose: empty keys read as
# "fill these in", and agents duly stopped to ask the user for a credential
# they don't need. Reading org data is what Organizations + the data_query
# tool are for, and that path carries its own short-lived credential.
DEFAULT_ENV_TEMPLATE = """\
# VibeFoundry secrets. Git-ignored — never commit or share this file.
#
# You do NOT need a credential here to explore or query your organization's
# data — connect through Organizations in the IDE and ask your question.
#
# An App Credential belongs here only for an app you publish, because a
# published app runs when you are not signed in. Mint one in the portal's
# App Credentials tab and add VF_APP_ID / VF_APP_KEY beside the gateway.

# Citizen Engineering Portal data gateway
VF_GATEWAY=https://data-gateway-625603147835.us-central1.run.app
"""


def _ensure_gitignored(project_folder: Path, entry: str) -> None:
    """Add `entry` to .gitignore if it isn't already listed. Build only writes
    a fresh .gitignore when it runs `git init` itself, so a project that was
    already a repo would otherwise have nothing ignoring the secrets file."""
    path = project_folder / ".gitignore"
    try:
        existing = path.read_text(encoding="utf-8") if path.exists() else ""
        if any(line.strip() == entry for line in existing.splitlines()):
            return
        separator = "" if (not existing or existing.endswith("\n")) else "\n"
        path.write_text(f"{existing}{separator}{entry}\n", encoding="utf-8")
    except OSError as e:
        print(f"[Build] could not add {entry} to .gitignore: {e}")


async def _fetch_remote_agents_md() -> Optional[bytes]:
    """The cascaded IDE-Agents/AGENTS.md — the authenticated proxy when signed
    in, else the public website path. One fetcher shared by /api/build,
    /api/track0/scaffold and /api/rules so the fallback chain cannot drift
    between them. (The website-download variant lives in Templates-Agents/ and
    is only bundled into downloadable zips.)"""
    stored = _read_stored_token()
    jwt = stored["token"] if stored else ""
    if jwt:
        try:
            async with httpx.AsyncClient(timeout=15.0) as client:
                res = await client.get(
                    f"{PROXY_BASE_URL}/IDE-Agents/AGENTS.md",
                    headers={
                        "Authorization": f"Bearer {jwt}",
                        "Accept": "application/vnd.github.raw",
                    },
                )
                if res.status_code == 200:
                    return res.content
        except Exception as e:
            print(f"[Build] Authenticated AGENTS.md fetch failed ({e}); falling back to public path")

    try:
        async with httpx.AsyncClient(timeout=10.0) as client:
            res = await client.get(f"{PUBLIC_TEMPLATE_FALLBACK_URL}/IDE-Agents/AGENTS.md")
            if res.status_code == 200:
                return res.content
    except Exception as e:
        print(f"[Build] Public AGENTS.md fallback also failed: {e}")
    return None


async def _ensure_agents_md(project_folder: Path, overwrite: bool) -> str:
    """AGENTS.md plus the CLAUDE.md shim in `project_folder`.

    Build refreshes AGENTS.md on every run (overwrite=True) so a rules change
    ships without a package release; scaffold only fills a gap, so a user who
    edited their rulebook keeps it. CLAUDE.md is never overwritten either way —
    Claude Code auto-loads it, not AGENTS.md (the Codex convention), and the
    one-line import makes both assistants read the same rules.

    Returns "written" | "present" | "unavailable"."""
    agents_dest = project_folder / "AGENTS.md"
    result = "present" if agents_dest.exists() else "unavailable"
    if overwrite or not agents_dest.exists():
        body = await _fetch_remote_agents_md()
        if body is not None:
            agents_dest.write_bytes(body)
            result = "written"
            # A rewritten rulebook invalidates the cached Track 0 section, or
            # /api/rules keeps serving the copy this build just replaced.
            _rules_cache.pop(str(project_folder), None)

    claude_dest = project_folder / "CLAUDE.md"
    # Only shim to a rulebook that is actually on disk. Writing `@AGENTS.md`
    # after a failed fetch leaves a dangling import on a fresh offline project,
    # and because CLAUDE.md is never overwritten the broken shim would survive
    # the first build that succeeds.
    if agents_dest.exists() and not claude_dest.exists():
        claude_dest.write_text("@AGENTS.md\n", encoding="utf-8")
    return result


@app.post("/api/build")
async def build_project():
    """Build the project structure — creates the input/output/app folders,
    fetches AGENTS.md, and drops a `.env` at the root for data-access secrets.
    Templates are no longer cascaded here; the user picks them explicitly via
    /api/templates/download.
    """
    if not state.project_folder:
        raise HTTPException(status_code=400, detail="No project folder selected")

    # Create folder structure (input_folder/, output_folder/, app_folder/, etc.)
    folders = setup_project_structure(state.project_folder)

    # jwt is still read here because env.template below needs it too.
    stored = _read_stored_token()
    jwt = stored["token"] if stored else ""
    await _ensure_agents_md(state.project_folder, overwrite=True)

    # Drop a `.env` at the project root — the one place data-access secrets
    # live. Never overwrite an existing one: it holds the user's pasted keys.
    env_dest = state.project_folder / ".env"
    env_created = False
    if not env_dest.exists():
        env_body = DEFAULT_ENV_TEMPLATE
        if jwt:
            try:
                async with httpx.AsyncClient(timeout=15.0) as client:
                    res = await client.get(
                        f"{PROXY_BASE_URL}/IDE-Agents/env.template",
                        headers={
                            "Authorization": f"Bearer {jwt}",
                            "Accept": "application/vnd.github.raw",
                        },
                    )
                    if res.status_code == 200 and res.text.strip():
                        env_body = res.text
            except Exception as e:
                print(f"[Build] env.template fetch failed ({e}); using built-in default")
        env_dest.write_text(env_body, encoding="utf-8")
        env_created = True

    # Initialize git repo if not already one
    git_initialized = False
    git_dir = state.project_folder / ".git"
    if not git_dir.exists():
        import subprocess
        try:
            subprocess.run(
                ["git", "init"],
                cwd=str(state.project_folder),
                capture_output=True,
                check=True
            )
            git_initialized = True

            # Create .gitignore with sensible defaults
            gitignore_path = state.project_folder / ".gitignore"
            if not gitignore_path.exists():
                gitignore_path.write_text(
                    "# Python\n"
                    "__pycache__/\n"
                    "*.py[cod]\n"
                    ".venv/\n"
                    "venv/\n"
                    "*.egg-info/\n"
                    "\n"
                    "# Node\n"
                    "node_modules/\n"
                    "\n"
                    "# Environment\n"
                    ".env\n"
                    ".env.local\n"
                    "\n"
                    "# OS\n"
                    ".DS_Store\n"
                    "Thumbs.db\n",
                    encoding="utf-8"
                )
        except (subprocess.CalledProcessError, FileNotFoundError):
            # git not installed or failed - continue without it
            pass

    _ensure_gitignored(state.project_folder, ".env")

    # Generate metadata now that folders exist
    generate_metadata(state.project_folder)

    # Restart watcher to pick up newly created folders
    if state.watcher:
        state.watcher.stop()
    state.watcher = FileWatcher(
        state.project_folder,
        on_data_change=notify_data_change,
        on_script_change=notify_script_change,
        on_output_file_change=notify_output_file_change
    )
    await state.watcher.start_async()

    return {
        "success": True,
        "folders": {k: str(v) for k, v in folders.items()},
        "agents_md_copied": (state.project_folder / "AGENTS.md").exists(),
        "env_created": env_created,
        "git_initialized": git_initialized
    }


# --- Track 0 rules -----------------------------------------------------------
#
# The Track 0 section of AGENTS.md, served on its own to a model that has
# never seen the rest of the file. That is why the section is written to be
# liftable verbatim, and why this endpoint can never fail: offline is not an
# excuse for the model not knowing where to put files.

BUILTIN_TRACK0_RULES = """\
# Track 0: Queries

A question about the data is answered by BUILDING: a small script per
question, run, and the answer read back out of the files it produced. Never
answer from memory, and never answer from rows seen mid-pipeline.

Folder, under the open project (`app_folder/` is created if absent):

    app_folder/scripts/{script_name}/
    |- app.py           orchestrator, runs the steps in order
    |- raw_pulls/       data pulled from the gateway or a public dataset
    |- steps/           step1_*.py, step2_*.py ... the processing logic
    `- final_output/    the answer: parquet / csv / xlsx / png / gif

`{script_name}` comes from the question (`georgia_top_accounts`). A follow-up
on the same subject reuses that folder and modifies its scripts.

Rules, all deliberate departures from Tracks 1-4:

- No `run_app.sh`, no `run_app.bat`, no `.command`, no `requirements.txt`.
  Nothing launches this standalone.
- No comments in any `.py`. None. No docstrings either. The chat answer is
  the explanation; these scripts are the working, not the document.
- Numbered steps live inside `steps/`. `app.py` is the only file at the root.
- Pulled data lands in `raw_pulls/`, not `input_folder/`.
- The answer lands in `final_output/`, not `output_folder/`. A merged table
  always; images too when the question asked for them.

Then run the script, read `final_output/`, and answer from that — a preview
inline, or the whole table when it is small enough to be worth showing. The
file and the answer cannot disagree, because they are the same thing.
"""

# Resolved rules keyed by project folder. The contract asks for a
# life-of-the-process cache; keying it means switching projects switches
# rulebooks, instead of serving project A's customised AGENTS.md to project B.
_rules_cache: dict[str, dict] = {}


def _extract_track0_section(markdown: str) -> Optional[str]:
    """From the line matching `^# Track 0` up to (not including) the next line
    matching `^# Track `, or EOF. Returns None when there is no Track 0
    section — an older AGENTS.md has to fall through to the next source rather
    than serve nothing."""
    lines = (markdown or "").splitlines()
    start = next((i for i, line in enumerate(lines) if line.startswith("# Track 0")), None)
    if start is None:
        return None
    end = next(
        (j for j in range(start + 1, len(lines)) if lines[j].startswith("# Track ")),
        len(lines),
    )
    section = "\n".join(lines[start:end]).strip()
    return section or None


@app.get("/api/rules")
async def get_rules():
    """The Track 0 rules as markdown, resolved project → remote → built-in.

    The project's own AGENTS.md wins: a user who customised their rulebook
    gets their rules, not ours."""
    key = str(state.project_folder or "")
    cached = _rules_cache.get(key)
    if cached:
        return cached

    section = None
    source = "builtin"

    if state.project_folder:
        local = state.project_folder / "AGENTS.md"
        try:
            if local.exists():
                section = _extract_track0_section(local.read_text(encoding="utf-8", errors="replace"))
                if section:
                    source = "project"
        except OSError as e:
            print(f"[Rules] could not read the project's AGENTS.md: {e}")

    if not section:
        body = await _fetch_remote_agents_md()
        if body:
            section = _extract_track0_section(body.decode("utf-8", errors="replace"))
            if section:
                source = "remote"

    if not section:
        section = BUILTIN_TRACK0_RULES.strip()
        source = "builtin"

    resolved = {"source": source, "markdown": section, "bytes": len(section.encode("utf-8"))}
    _rules_cache[key] = resolved
    return resolved


# --- Template selective download / delete -----------------------------------


def _require_jwt() -> str:
    stored = _read_stored_token()
    jwt = stored["token"] if stored else ""
    if not jwt:
        raise HTTPException(status_code=401, detail="Not signed in")
    return jwt


def _restart_watcher() -> None:
    if not state.project_folder:
        return
    if state.watcher:
        state.watcher.stop()
    state.watcher = FileWatcher(
        state.project_folder,
        on_data_change=notify_data_change,
        on_script_change=notify_script_change,
        on_output_file_change=notify_output_file_change,
    )


@app.get("/api/templates/catalog")
async def get_templates_catalog():
    """Fetch catalog.json live from the proxy. Not written to disk — the
    picker UI reads it on every open so new templates appear without an IDE
    update."""
    jwt = _require_jwt()
    timeout = httpx.Timeout(connect=10.0, read=30.0, write=10.0, pool=10.0)
    try:
        async with httpx.AsyncClient(timeout=timeout) as client:
            res = await client.get(
                f"{PROXY_BASE_URL}/catalog.json",
                headers={
                    "Authorization": f"Bearer {jwt}",
                    "Accept": "application/vnd.github.raw",
                },
            )
            res.raise_for_status()
            return res.json()
    except httpx.HTTPStatusError as e:
        raise HTTPException(status_code=e.response.status_code, detail=f"Proxy error fetching catalog: {e}")
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Could not fetch catalog: {e}")


_ICON_MIME = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".svg": "image/svg+xml",
    ".webp": "image/webp",
}


@app.get("/api/templates/icon/{name}")
async def get_template_icon(name: str):
    """Stream a template icon from the proxy. JWT-authenticated; no disk write."""
    if not name or "/" in name or "\\" in name or ".." in name:
        raise HTTPException(status_code=400, detail="Invalid icon name")
    ext = Path(name).suffix.lower()
    if ext not in _ICON_MIME:
        raise HTTPException(status_code=400, detail="Unsupported icon type")
    jwt = _require_jwt()
    timeout = httpx.Timeout(connect=10.0, read=30.0, write=10.0, pool=10.0)
    try:
        async with httpx.AsyncClient(timeout=timeout) as client:
            res = await client.get(
                f"{PROXY_BASE_URL}/{name}",
                headers={
                    "Authorization": f"Bearer {jwt}",
                    "Accept": "application/vnd.github.raw",
                },
            )
            res.raise_for_status()
            return Response(content=res.content, media_type=_ICON_MIME[ext])
    except httpx.HTTPStatusError as e:
        raise HTTPException(status_code=e.response.status_code, detail=f"Icon not found: {name}")
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Could not fetch icon: {e}")


class TemplateDownloadRequest(BaseModel):
    template_id: str


def _valid_template_id(tid: str) -> bool:
    if not tid:
        return False
    return all(c.islower() or c.isdigit() or c == "_" for c in tid)


@app.post("/api/templates/download")
async def download_template(req: TemplateDownloadRequest):
    """Download one template subtree into {project}/templates/{template_id}/.
    Repeated downloads overwrite that subfolder; other downloaded templates
    are untouched."""
    if not state.project_folder:
        raise HTTPException(status_code=400, detail="No project folder selected")
    if not _valid_template_id(req.template_id):
        raise HTTPException(status_code=400, detail="Invalid template_id")
    jwt = _require_jwt()

    templates_root = state.project_folder / "templates"
    templates_root.mkdir(parents=True, exist_ok=True)

    try:
        written = await _cascade_templates_via_proxy(
            templates_root,
            jwt,
            subpath=req.template_id,
        )
    except httpx.HTTPStatusError as e:
        raise HTTPException(status_code=e.response.status_code, detail=f"Proxy error: {e}")
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Download failed: {e}")

    generate_metadata(state.project_folder)
    _restart_watcher()
    if state.watcher:
        await state.watcher.start_async()

    return {
        "success": True,
        "template_id": req.template_id,
        "files_cascaded": len(written),
    }


@app.delete("/api/templates")
async def delete_templates():
    """Remove the entire templates/ folder from the current project."""
    if not state.project_folder:
        raise HTTPException(status_code=400, detail="No project folder selected")

    templates_dir = state.project_folder / "templates"
    try:
        templates_dir.resolve().relative_to(state.project_folder.resolve())
    except ValueError:
        raise HTTPException(status_code=403, detail="Access denied")

    deleted = False
    if templates_dir.exists():
        shutil.rmtree(templates_dir)
        deleted = True

    generate_metadata(state.project_folder)
    _restart_watcher()
    if state.watcher:
        await state.watcher.start_async()

    return {"success": True, "deleted": deleted}


# --- Public data library ------------------------------------------------------
#
# Static datasets served straight off vibefoundry.ai. No identity, no token,
# no gate — they are public files. Downloads land in {project}/input_folder/,
# created on demand, because that is where scripts already look for inputs.

PUBLIC_DATA_BASE_URL = "https://vibefoundry.ai/public_data"

_DATA_TIMEOUT = httpx.Timeout(connect=10.0, read=180.0, write=30.0, pool=10.0)


class DataDownloadRequest(BaseModel):
    dataset_id: str


def _valid_dataset_id(did: str) -> bool:
    """Same shape rule as template ids. This value is interpolated into an
    upstream URL and used as a filename stem, so anything outside
    [a-z0-9_] is rejected rather than escaped."""
    if not did or len(did) > 128:
        return False
    return all(c.islower() or c.isdigit() or c == "_" for c in did)


def _input_folder() -> Path:
    """The project's input_folder, created if it isn't there yet."""
    if not state.project_folder:
        raise HTTPException(status_code=400, detail="No project folder selected")
    folder = state.project_folder / "input_folder"
    folder.mkdir(parents=True, exist_ok=True)
    return folder


def _safe_dest(folder: Path, filename: str) -> Path:
    """Resolve `filename` inside `folder`, refusing anything that escapes it.
    The name comes from a remote manifest, so it is treated as untrusted."""
    name = Path(str(filename or "")).name
    if not name or name in (".", ".."):
        raise HTTPException(status_code=502, detail="Upstream sent an unusable filename")
    dest = (folder / name).resolve()
    try:
        dest.relative_to(folder.resolve())
    except ValueError:
        raise HTTPException(status_code=502, detail="Upstream filename escapes the destination folder")
    return dest


async def _public_dataset_meta(client: httpx.AsyncClient, dataset_id: str) -> dict:
    """The per-dataset json, which is what knows the parquet's filename; the
    manifest only carries display metadata.

    A single-page app answers unknown paths with index.html and a 200, so
    raise_for_status() cannot tell us the id was bogus — only failing to parse
    JSON can. Treat that as "no such dataset" rather than letting a decode
    error surface as a 502."""
    res = await client.get(f"{PUBLIC_DATA_BASE_URL}/{dataset_id}.json")
    res.raise_for_status()
    try:
        return res.json()
    except ValueError:
        raise HTTPException(status_code=404, detail=f"No public dataset named '{dataset_id}'")


async def _after_download() -> None:
    """Refresh metadata and the file watcher so the new file shows up in the
    tree immediately, the same way a template download does. _restart_watcher
    only builds the watcher — it has to be started too, or the tree goes quiet
    after the first download."""
    generate_metadata(state.project_folder)
    _restart_watcher()
    if state.watcher:
        await state.watcher.start_async()


@app.get("/api/data/public/catalog")
async def get_public_data_catalog():
    """The public dataset list, read live so a dataset added by
    refresh_public_data.py appears here without an IDE release.

    The manifest decides which datasets exist and in what order; the
    per-dataset json carries the description, column documentation and row
    preview. Both are fetched, exactly as the website's Public Data page does,
    so the IDE shows the same thing rather than a thinner summary of it."""
    try:
        async with httpx.AsyncClient(timeout=_DATA_TIMEOUT) as client:
            res = await client.get(f"{PUBLIC_DATA_BASE_URL}/manifest.json")
            res.raise_for_status()
            manifest = res.json()

            ids = [d.get("id") for d in manifest.get("datasets", []) if d.get("id")]
            # Gathered concurrently, and one missing dataset must not blank the
            # whole library — same tolerance the website page has.
            async def fetch_one(did: str):
                try:
                    r = await client.get(f"{PUBLIC_DATA_BASE_URL}/{did}.json")
                    r.raise_for_status()
                    body = r.json()
                    return body if isinstance(body, dict) else None
                except Exception:
                    return None

            details = await asyncio.gather(*(fetch_one(i) for i in ids))
    except httpx.HTTPStatusError as e:
        raise HTTPException(status_code=e.response.status_code, detail=f"Could not fetch catalog: {e}")
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Could not fetch catalog: {e}")

    by_id = {d["id"]: d for d in details if d and d.get("id")}
    merged = []
    for entry in manifest.get("datasets", []):
        did = entry.get("id")
        # Manifest fields win for the summary numbers; the detail json supplies
        # description, columns and preview.
        merged.append({**by_id.get(did, {}), **entry})

    return {"datasets": merged, "generatedAt": manifest.get("generatedAt")}


@app.get("/api/data/public/file/{dataset_id}")
async def get_public_data_file(dataset_id: str):
    """Stream one public parquet through the IDE.

    The filter UI reads the whole file in the browser. Fetching it straight
    from vibefoundry.ai would be cross-origin from the IDE's localhost, so it
    comes through here instead and is same-origin to the pane. Also spares us
    depending on range requests and CORS headers on the host."""
    if not _valid_dataset_id(dataset_id):
        raise HTTPException(status_code=400, detail="Invalid dataset_id")
    try:
        async with httpx.AsyncClient(timeout=_DATA_TIMEOUT, follow_redirects=True) as client:
            meta_res = await client.get(f"{PUBLIC_DATA_BASE_URL}/{dataset_id}.json")
            meta_res.raise_for_status()
            try:
                meta = meta_res.json()
            except ValueError:
                raise HTTPException(status_code=404, detail=f"No public dataset named '{dataset_id}'")
            source_file = meta.get("sourceFile") or f"{dataset_id}.parquet"
            url = meta.get("downloadUrl") or f"{PUBLIC_DATA_BASE_URL}/{source_file}"
            file_res = await client.get(url)
            file_res.raise_for_status()
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Could not read dataset: {e}")

    return Response(
        content=file_res.content,
        media_type="application/vnd.apache.parquet",
        headers={"x-vf-filename": source_file},
    )


@app.post("/api/data/public/save-cut")
async def save_public_data_cut(
    file: UploadFile = File(...),
    filename: str = Form(...),
):
    """Write a filtered cut, built in the browser, into input_folder/.

    The full-file path downloads server-side; a cut only exists client-side
    after filtering, so it is posted back here rather than re-derived."""
    folder = _input_folder()
    dest = _safe_dest(folder, filename)
    dest.write_bytes(await file.read())
    await _after_download()
    return {"success": True, "filename": dest.name, "bytes": dest.stat().st_size}


@app.post("/api/data/public/download")
async def download_public_data(req: DataDownloadRequest):
    """Fetch one public dataset's parquet into {project}/input_folder/."""
    if not _valid_dataset_id(req.dataset_id):
        raise HTTPException(status_code=400, detail="Invalid dataset_id")
    folder = _input_folder()

    try:
        async with httpx.AsyncClient(timeout=_DATA_TIMEOUT, follow_redirects=True) as client:
            meta = await _public_dataset_meta(client, req.dataset_id)

            source_file = meta.get("sourceFile") or f"{req.dataset_id}.parquet"
            url = meta.get("downloadUrl") or f"{PUBLIC_DATA_BASE_URL}/{source_file}"

            dest = _safe_dest(folder, source_file)
            file_res = await client.get(url)
            file_res.raise_for_status()
            dest.write_bytes(file_res.content)
    except HTTPException:
        raise
    except httpx.HTTPStatusError as e:
        raise HTTPException(status_code=e.response.status_code, detail=f"Download failed: {e}")
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Download failed: {e}")

    await _after_download()
    return {
        "success": True,
        "dataset_id": req.dataset_id,
        "filename": dest.name,
        "bytes": dest.stat().st_size,
    }


# --- Track 0 scaffold --------------------------------------------------------
#
# Track 0 answers by building, into a folder shape that is deliberately not
# Tracks 1-4's: pulls land in raw_pulls/ instead of input_folder/, the answer
# in final_output/ instead of output_folder/, and nothing launches standalone.
# /api/rules serves the rules themselves.

TRACK0_SUBFOLDERS = ("raw_pulls", "steps", "final_output")


def _valid_script_name(name: str) -> bool:
    """The model picks this name out of the question, so it is untrusted input
    that becomes a folder name: allow-listed, not sanitized."""
    if not name or len(name) > 96:
        return False
    return all(c.isalnum() or c in ("_", "-") for c in name)


def _is_track0_folder(folder: Path) -> bool:
    """Whether an existing folder is one Track 0 may take over: one it already
    owns (it has at least one of the three subfolders) or an empty one.
    Anything else is a Track 1-4 task folder that happens to share the name,
    and that is someone's app — see _script_folders."""
    if not folder.is_dir():
        return False
    if any((folder / sub).is_dir() for sub in TRACK0_SUBFOLDERS):
        return True
    return not any(folder.iterdir())


def _script_folders(script_name: str) -> dict:
    """The Track 0 folder for `script_name` plus its three subfolders, created
    if missing. `reused` reports whether the folder was already there — a
    follow-up on the same subject is meant to land back in the same one.

    The model picks this name out of free text, so it can collide with an
    existing Track 1-4 task folder. Reuse is only silent for a folder Track 0
    already owns; injecting raw_pulls/, steps/ and final_output/ into someone's
    app is a conflict the caller has to resolve by picking another name."""
    if not state.project_folder:
        raise HTTPException(status_code=400, detail="No project folder selected")
    if not _valid_script_name(script_name):
        raise HTTPException(status_code=400, detail="script_name must be letters, digits, _ or -")

    folder = state.project_folder / "app_folder" / "scripts" / script_name
    reused = folder.exists()
    if reused and not _is_track0_folder(folder):
        raise HTTPException(
            status_code=409,
            detail=(
                f"app_folder/scripts/{script_name} already exists and is not a Track 0 "
                "folder. Pick a different script_name."
            ),
        )
    paths = {"folder": folder, "reused": reused}
    for sub in TRACK0_SUBFOLDERS:
        # parents=True also creates app_folder/ and app_folder/scripts/ for a
        # project that has never been built.
        path = folder / sub
        path.mkdir(parents=True, exist_ok=True)
        paths[sub] = path
    return paths


class Track0ScaffoldRequest(BaseModel):
    script_name: str


@app.post("/api/track0/scaffold")
async def track0_scaffold(req: Track0ScaffoldRequest):
    """Create (or silently reuse) the Track 0 folder for one question, and make
    sure the project has the rulebook the model is about to write against."""
    paths = _script_folders(req.script_name)
    agents_md = await _ensure_agents_md(state.project_folder, overwrite=False)
    await _after_download()
    return {
        "script_name": req.script_name,
        "folder": str(paths["folder"]),
        "raw_pulls": str(paths["raw_pulls"]),
        "steps": str(paths["steps"]),
        "final_output": str(paths["final_output"]),
        "reused": paths["reused"],
        "agents_md": agents_md,
    }


# --- Organizations: connect, catalogue, query --------------------------------
#
# The IDE talks straight to the client's own hub and gateway. Nothing in this
# section touches vibefoundry.ai: the hub mints a short-lived personal
# credential, the gateway answers SQL against it, and the credential is stored
# only on this machine. That is the sovereignty guarantee, and it is why the
# org list is bundled (see organizations.py) rather than looked up remotely.
#
# The credential itself never leaves this file's helpers — /api/org/status,
# /api/org/list and /api/health return connection facts, never a key.

_ORG_STATE_TTL_SECONDS = 600  # 10 minutes, same as the sign-in nonce
_pending_org_states: dict[str, dict] = {}  # nonce -> {"hub_url": str, "expires": ts}

# One in-flight connect per org. A pane catalogue refresh and a plugin tool
# call hitting the same expired credential are two callers with one intent:
# without this they open two browser tabs and mint two nonces, and the user
# signs in twice. org_id (or the hub URL for a pasted hub) -> the handshake
# already running for it.
_org_connect_inflight: dict[str, dict] = {}  # key -> {"nonce", "expires", "result"}

_ORG_TIMEOUT = httpx.Timeout(connect=10.0, read=180.0, write=30.0, pool=10.0)

# Rows returned inline by /api/org/query. Anything past this spills to a
# parquet in the script's final_output/ — an LLM caller cannot read 200k rows,
# and a pane cannot render them, but a step downstream can read the file.
QUERY_PREVIEW_ROWS = 500

# Where a spill lands when the caller named no script. Track 0's answer always
# lives in a final_output/, so the fallback is a real Track 0 folder rather
# than a stray file somewhere else.
QUERY_FALLBACK_SCRIPT = "queries"


class OrgReauthRequired(Exception):
    """The org's stored credential is gone or was refused. Raised deep in a
    gateway call so every endpoint can answer with the reconnect handshake
    instead of an error the UI would have to special-case."""

    def __init__(self, org_id: str, hub_url: Optional[str] = None):
        super().__init__(org_id)
        self.org_id = org_id
        # Captured before the credential is dropped: an org reached by a pasted
        # hub URL is not in the bundled list, so once its entry is gone there
        # would be no hub left to send the user back to.
        self.hub_url = hub_url


class OrgNotConnected(Exception):
    """This org has no stored credential at all — never connected on this
    machine, or disconnected. Distinct from OrgReauthRequired on purpose:
    auto-reauth is scoped to a gateway 401/403 on a credential that exists, so
    this case is reported rather than self-healed. Opening a hub sign-in tab
    for an org the user never connected is a surprise, and the caller's remedy
    is one connect_organization call."""

    def __init__(self, org_id: str):
        super().__init__(org_id)
        self.org_id = org_id


def _not_connected_response(org_id: str) -> dict:
    """The instant "call connect_organization" answer. Same wire shape the
    plugin already handles for a lapsed credential, minus the browser."""
    return {"status": "reauth_required", "org_id": org_id}


def _orgs_store_path() -> Path:
    home = Path.home() / ".vibefoundry"
    home.mkdir(parents=True, exist_ok=True)
    return home / "orgs.json"


def _parse_iso_ts(value) -> Optional[float]:
    """ISO-8601 → unix seconds, or None if it can't be read. The hub mints
    `expires` with a trailing 'Z', which fromisoformat rejects before 3.11,
    so swap it for the offset it stands for."""
    if not value or not isinstance(value, str):
        return None
    text = value.strip()
    if text.endswith("Z"):
        text = text[:-1] + "+00:00"
    try:
        parsed = datetime.fromisoformat(text)
    except ValueError:
        return None
    if parsed.tzinfo is None:
        parsed = parsed.replace(tzinfo=timezone.utc)
    return parsed.timestamp()


def _write_orgs(orgs: dict) -> None:
    """Persist the org store 0600. The mode is set at creation rather than
    chmod'ed afterwards, because the file holds live gateway keys and even a
    brief world-readable window is a leak."""
    path = _orgs_store_path()
    fd = os.open(str(path), os.O_WRONLY | os.O_CREAT | os.O_TRUNC, 0o600)
    with os.fdopen(fd, "w", encoding="utf-8") as handle:
        handle.write(json.dumps({"orgs": orgs}, indent=2))
    try:
        os.chmod(path, 0o600)  # a file that already existed keeps its old mode
    except OSError:
        pass


def _read_orgs() -> dict:
    """Stored org credentials keyed by org_id, with expired ones dropped.

    Pruning happens on read so an expired credential never reaches a request:
    the endpoint sees "not connected" and asks for a reconnect, instead of
    spending a round trip discovering the gateway agrees."""
    path = _orgs_store_path()
    if not path.exists():
        return {}
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return {}
    orgs = data.get("orgs")
    if not isinstance(orgs, dict):
        return {}

    now = time.time()
    live = {}
    for org_id, entry in orgs.items():
        if not isinstance(entry, dict) or not entry.get("app_id") or not entry.get("app_key"):
            continue
        expires_at = _parse_iso_ts(entry.get("expires"))
        # An unreadable or absent expiry is not treated as expired — the
        # gateway is the authority, and a 401 there drives the reconnect.
        if expires_at is not None and expires_at <= now:
            continue
        live[org_id] = entry

    if len(live) != len(orgs):
        _write_orgs(live)
    return live


def _org_entry(org_id: str) -> Optional[dict]:
    return _read_orgs().get(str(org_id or "").strip())


def _drop_org(org_id: str) -> bool:
    orgs = _read_orgs()
    if org_id not in orgs:
        return False
    orgs.pop(org_id, None)
    _write_orgs(orgs)
    return True


def _org_public_view(entry: dict) -> dict:
    """Connection facts for one org. Built field by field on purpose: a
    dict-minus-the-key would leak any secret field added here later."""
    expires_at = _parse_iso_ts(entry.get("expires"))
    return {
        "org_id": entry.get("org_id"),
        "org_name": entry.get("org_name"),
        "hub_url": entry.get("hub_url"),
        "gateway": entry.get("gateway"),
        "email": entry.get("email"),
        "expires": entry.get("expires"),
        "seconds_to_expiry": (
            max(0, int(expires_at - time.time())) if expires_at is not None else None
        ),
        "connected_at": entry.get("connected_at"),
        "tables": entry.get("tables"),
    }


def _purge_expired_org_states() -> None:
    now = time.time()
    for nonce in [n for n, v in _pending_org_states.items() if v.get("expires", 0) < now]:
        _pending_org_states.pop(nonce, None)
    for key in [k for k, v in _org_connect_inflight.items() if v.get("expires", 0) < now]:
        _org_connect_inflight.pop(key, None)


def _clear_org_connect_inflight(nonce: str) -> None:
    """Drop the in-flight record whose nonce just came back through the
    callback: that handshake is over, so the next expiry may open a tab again."""
    for key in [k for k, v in _org_connect_inflight.items() if v.get("nonce") == nonce]:
        _org_connect_inflight.pop(key, None)


def _valid_table_id(tid: str) -> bool:
    """Wider than _valid_dataset_id: an org's catalogue ids are the client's
    to choose, so capitals and hyphens are allowed. Still an allow-list — the
    value is interpolated into a gateway URL path."""
    if not tid or len(tid) > 128:
        return False
    return all(c.isalnum() or c in ("_", "-") for c in tid)


async def _gateway_request(entry: dict, method: str, path: str, json_body: Optional[dict] = None) -> httpx.Response:
    """One authenticated call to an org's gateway.

    A 401/403 means the personal credential expired or was revoked, so the
    stored copy is dropped here rather than left to fail again on the next
    call, and the caller is told to send the user back through /api/org/connect."""
    url = str(entry.get("gateway", "")).rstrip("/") + path
    headers = {"Authorization": f"Bearer {entry['app_id']}.{entry['app_key']}"}
    try:
        async with httpx.AsyncClient(timeout=_ORG_TIMEOUT, follow_redirects=True) as client:
            res = await client.request(method, url, headers=headers, json=json_body)
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Could not reach {entry.get('org_name') or entry.get('org_id')}: {e}")

    if res.status_code in (401, 403):
        _drop_org(entry.get("org_id"))
        raise OrgReauthRequired(entry.get("org_id"), entry.get("hub_url"))
    return res


def _raise_for_gateway_error(res: httpx.Response) -> None:
    """Re-raise a gateway refusal as our own, carrying its message through.
    401/403 never reach here — _gateway_request turns those into a reconnect."""
    if res.status_code < 400:
        return
    try:
        detail = res.json().get("detail", "")
    except Exception:
        detail = res.text[:500]
    raise HTTPException(status_code=res.status_code, detail=detail or f"Gateway returned {res.status_code}")


def _gateway_json(res: httpx.Response) -> dict:
    _raise_for_gateway_error(res)
    try:
        return res.json()
    except ValueError:
        raise HTTPException(status_code=502, detail="Gateway returned a non-JSON response")


def _require_org(org_id: str) -> dict:
    entry = _org_entry(org_id)
    if not entry:
        # No stored credential is not a refused one: nothing to re-authenticate,
        # so say so instead of opening a browser the caller did not ask for.
        raise OrgNotConnected(org_id)
    return entry


def _callback_host(request: Optional[Request]) -> str:
    """The Host header, not a hardcoded port: the backend binds whatever port
    was free, and the hub only accepts a loopback callback."""
    if request is None:
        return "127.0.0.1:8765"
    return request.headers.get("host", "127.0.0.1:8765")


async def _start_org_connect(
    hub_url: str,
    host: str,
    key: Optional[str] = None,
    dedupe: bool = False,
) -> dict:
    """Mint a nonce and open the hub's /connect in the user's browser.

    Unlike /api/auth/start, which hands the URL back for the frontend to open,
    this process opens the browser itself: connect is driven from a plugin tool
    call as often as from the UI, and a sandboxed pane cannot open a top-level
    window. The URL is still returned so the UI can offer it as a link when
    webbrowser has nothing to launch.

    `dedupe` is for the automatic path. A pane catalogue refresh and a plugin
    tool call hitting the same expired credential are two callers with one
    intent: without the guard they open two tabs and mint two nonces, and the
    user signs in twice. An explicit /api/org/connect is a deliberate click, so
    it always opens a tab — but it takes over the slot, so an auto-reauth
    arriving behind it rides on the handshake already in the browser."""
    _purge_expired_org_states()
    slot = key or hub_url
    if dedupe:
        running = _org_connect_inflight.get(slot)
        if running:
            return dict(running["result"])

    nonce = secrets.token_urlsafe(32)
    _pending_org_states[nonce] = {"hub_url": hub_url, "expires": time.time() + _ORG_STATE_TTL_SECONDS}

    callback_url = f"http://{host}/org/callback"
    connect_url = (
        f"{hub_url}/connect"
        f"?state={nonce}"
        f"&callback={urllib.parse.quote(callback_url, safe='')}"
    )
    opened = await asyncio.to_thread(webbrowser.open, connect_url)
    result = {"hub_url": hub_url, "url": connect_url, "browser_opened": bool(opened)}
    _org_connect_inflight[slot] = {
        "nonce": nonce,
        "expires": time.time() + _ORG_STATE_TTL_SECONDS,
        "result": result,
    }
    state.open_org_catalog = True
    return dict(result)


def _hub_for_org(org_id: str, hub_url: Optional[str]) -> Optional[str]:
    """Where to send the user back to. The exception's copy first, because the
    stored entry may already have been dropped by the 403 that got us here."""
    if hub_url:
        return hub_url
    entry = _org_entry(org_id)
    if entry and entry.get("hub_url"):
        return entry["hub_url"]
    org = find_organization(org_id)
    return org["hub_url"] if org else None


async def _reauth_response(org_id: str, hub_url: Optional[str] = None, host: str = "127.0.0.1:8765") -> dict:
    """Re-authenticate instead of reporting. Merely returning "reconnect" costs
    a model turn, and with a one-hour credential that is an hourly stumble — so
    the credential is cleared and the browser opened here, and the caller polls
    /api/org/status and retries its call once."""
    hub = _hub_for_org(org_id, hub_url)
    _drop_org(org_id)
    if not hub:
        # A pasted hub whose entry is already gone: nothing to reopen.
        return {"status": "reauth_required", "org_id": org_id}
    started = await _start_org_connect(hub, host, key=org_id or hub, dedupe=True)
    return {"status": "reauth_started", "org_id": org_id, **started}


# --- SQL guard rails ---------------------------------------------------------
#
# The gateway enforces its own copy of these rules and is the authority for
# org queries. This copy exists for the public path, where THIS process is the
# one executing the SQL: without it, `SELECT * FROM read_parquet('~/.ssh/id_rsa')`
# would be a working local file read dressed up as a data question.

_SQL_BANNED = re.compile(
    r"\b(read_parquet|read_csv|read_ndjson|scan_\w+|COPY|ATTACH|INSERT|UPDATE"
    r"|DELETE|CREATE|DROP|ALTER|TRUNCATE|GRANT|EXECUTE)\b",
    re.IGNORECASE,
)
_SQL_TABLE_REF = re.compile(r'\b(?:FROM|JOIN)\s+"?([A-Za-z_][A-Za-z0-9_\-]*)"?', re.IGNORECASE)


def _validate_local_sql(sql: str) -> str:
    """Return the single SELECT statement to execute, or raise 400."""
    text = (sql or "").strip()
    if not text:
        raise HTTPException(status_code=400, detail="No SQL supplied")
    # Comments are rejected rather than stripped: stripping is where smuggling
    # gets in, because the stripper and the parser never agree on the edges.
    if "--" in text or "/*" in text:
        raise HTTPException(status_code=400, detail="SQL comments are not allowed")
    if text.endswith(";"):
        text = text[:-1].rstrip()
    if ";" in text:
        raise HTTPException(status_code=400, detail="Only one statement per query")
    if not re.match(r"^(SELECT|WITH)\b", text, re.IGNORECASE):
        raise HTTPException(status_code=400, detail="Query must start with SELECT or WITH")
    banned = _SQL_BANNED.search(text)
    if banned:
        raise HTTPException(status_code=400, detail=f"'{banned.group(1)}' is not allowed in a query")
    return text


def _referenced_tables(sql: str) -> list[str]:
    seen = []
    for match in _SQL_TABLE_REF.finditer(sql or ""):
        name = match.group(1)
        if name.lower() not in ("select", "lateral") and name not in seen:
            seen.append(name)
    return seen


def _json_scalar(v):
    """NaN/Inf are valid float() but not valid JSON; the frontend's fetch
    rejects the whole payload when one slips through."""
    if isinstance(v, float) and (math.isnan(v) or math.isinf(v)):
        return None
    return v


def _frame_to_rows(df: pl.DataFrame) -> tuple[list[str], list[list]]:
    """Columns + row lists, JSON-safe. polars' own writer is used for the
    conversion so dates, durations and decimals arrive as strings instead of
    objects the encoder chokes on."""
    columns = list(df.columns)
    if df.height == 0:
        return columns, []
    records = json.loads(df.write_json())
    return columns, [[_json_scalar(rec.get(c)) for c in columns] for rec in records]


def _query_result_stem(tables_used: list) -> str:
    base = ""
    if tables_used:
        base = "".join(c if (c.isalnum() or c in ("_", "-")) else "_" for c in str(tables_used[0]))
    return f"{base or 'query'}_{time.strftime('%Y%m%d_%H%M%S')}"


def _spill_result(df: pl.DataFrame, tables_used: list, script_name: Optional[str] = None) -> Optional[str]:
    """Write a too-big result into the script's final_output/ and return its
    path. Returns None when there is no project folder — the preview still
    answers the question, and losing the spill is better than failing the query."""
    if not state.project_folder:
        return None
    # No fallback here: /api/org/query resolves a caller-supplied script_name
    # before it runs any SQL, so by this point the name is known good. Callers
    # treat a raised failure as "lost the spill, keep the answer".
    folder = _script_folders(script_name or QUERY_FALLBACK_SCRIPT)["final_output"]
    dest = folder / f"{_query_result_stem(tables_used)}.parquet"
    df.write_parquet(dest)
    return str(dest)


async def _ensure_public_parquet(dataset_id: str, force: bool = False) -> tuple[Path, dict]:
    """The public dataset's parquet on disk in input_folder/, downloading it
    only if it isn't already there. Questions about public data are answered
    locally, so the file has to exist — but a second question about the same
    dataset must not re-download 4 MB to ask it."""
    if not _valid_dataset_id(dataset_id):
        raise HTTPException(status_code=400, detail="Invalid dataset_id")
    folder = _input_folder()
    async with httpx.AsyncClient(timeout=_DATA_TIMEOUT, follow_redirects=True) as client:
        meta = await _public_dataset_meta(client, dataset_id)
        source_file = meta.get("sourceFile") or f"{dataset_id}.parquet"
        dest = _safe_dest(folder, source_file)
        if force or not dest.exists():
            url = meta.get("downloadUrl") or f"{PUBLIC_DATA_BASE_URL}/{source_file}"
            file_res = await client.get(url)
            file_res.raise_for_status()
            dest.write_bytes(file_res.content)
    return dest, meta


async def _public_manifest() -> dict:
    try:
        async with httpx.AsyncClient(timeout=_DATA_TIMEOUT) as client:
            res = await client.get(f"{PUBLIC_DATA_BASE_URL}/manifest.json")
            res.raise_for_status()
            return res.json()
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Could not fetch public catalog: {e}")


async def _collect_public_sql(sql: str, limit: Optional[int]) -> tuple:
    """Execute SQL over public datasets locally with polars, returning
    (DataFrame, tables_used, elapsed_ms).

    Every table the SQL names is resolved against the public manifest and
    registered as a LazyFrame under its catalogue id, so the same SQL a user
    would send to a gateway works here unchanged."""
    statement = _validate_local_sql(sql)
    referenced = _referenced_tables(statement)
    manifest = await _public_manifest()
    known = {d.get("id") for d in manifest.get("datasets", []) if d.get("id")}
    wanted = [t for t in referenced if t in known]
    if not wanted:
        raise HTTPException(
            status_code=400,
            detail=(
                f"No public dataset matches {referenced or 'the tables in this query'}. "
                "Call /api/org/catalog for the ids that exist."
            ),
        )

    frames = {}
    for table_id in wanted:
        path, _ = await _ensure_public_parquet(table_id)
        frames[table_id] = pl.scan_parquet(path)

    started = time.time()
    try:
        lazy = pl.SQLContext(frames=frames).execute(statement)
        if limit and limit > 0:
            lazy = lazy.limit(int(limit))
        df = await asyncio.to_thread(lazy.collect)
    except HTTPException:
        raise
    except Exception as e:
        # Contract: a bad query is the caller's 400, never our 500.
        raise HTTPException(status_code=400, detail=str(e)[:500])
    return df, wanted, int((time.time() - started) * 1000)


async def _run_public_sql(sql: str, limit: Optional[int], script_name: Optional[str] = None) -> dict:
    df, wanted, elapsed_ms = await _collect_public_sql(sql, limit)
    columns, rows = _frame_to_rows(df.head(QUERY_PREVIEW_ROWS))
    spilled = None
    if df.height > QUERY_PREVIEW_ROWS:
        # Guarded exactly as the org path is: the answer is already computed,
        # and a failed write into final_output/ must not turn it into a 500.
        try:
            spilled = _spill_result(df, wanted, script_name)
        except Exception as e:
            print(f"[Public] could not spill query result: {e}")
    return {
        "status": "ok",
        "org_id": PUBLIC_ORG_ID,
        "columns": columns,
        "rows": rows,
        "row_count": df.height,
        "preview_rows": len(rows),
        "truncated": False,
        "tables_used": wanted,
        "elapsed_ms": elapsed_ms,
        "spilled_to": spilled,
    }


async def _run_org_sql(entry: dict, sql: str, limit: Optional[int], script_name: Optional[str] = None) -> dict:
    body = {"sql": sql, "format": "json"}
    if limit and limit > 0:
        body["limit"] = int(limit)
    # No client-side SQL validation on this path: the gateway enforces the
    # lock-down and is the authority, and a second, drifting copy of the rules
    # here would start rejecting queries the gateway is happy to run.
    res = await _gateway_request(entry, "POST", "/v1/query", json_body=body)
    payload = _gateway_json(res)

    columns = payload.get("columns") or []
    rows = payload.get("rows") or []
    row_count = payload.get("row_count", len(rows))
    spilled = None
    if len(rows) > QUERY_PREVIEW_ROWS:
        try:
            spilled = _spill_result(
                pl.DataFrame(rows, schema=columns, orient="row"),
                payload.get("tables_used") or [],
                script_name,
            )
        except Exception as e:
            print(f"[Org] could not spill query result: {e}")
        rows = rows[:QUERY_PREVIEW_ROWS]

    return {
        "status": "ok",
        "org_id": entry.get("org_id"),
        "columns": columns,
        "rows": [[_json_scalar(v) for v in row] for row in rows],
        "row_count": row_count,
        "preview_rows": len(rows),
        "truncated": bool(payload.get("truncated")),
        "tables_used": payload.get("tables_used") or [],
        "elapsed_ms": payload.get("elapsed_ms"),
        "spilled_to": spilled,
    }


# --- Organization endpoints --------------------------------------------------


class OrgConnectRequest(BaseModel):
    org_id: Optional[str] = None
    hub_url: Optional[str] = None


class OrgDisconnectRequest(BaseModel):
    org_id: str


class OrgQueryRequest(BaseModel):
    org_id: str
    sql: str
    limit: Optional[int] = None
    script_name: Optional[str] = None


class OrgPullRequest(BaseModel):
    org_id: str
    table_id: str = ""
    sql: Optional[str] = None
    filename: Optional[str] = None
    script_name: Optional[str] = None


@app.get("/api/org/list")
async def org_list():
    """The bundled organizations plus whether each is connected. Orgs reached
    by a pasted hub URL aren't in the bundle, so connected-but-unlisted ones
    are appended — otherwise they'd vanish from the picker after connecting."""
    connected = _read_orgs()
    listed = []
    seen = set()
    for org in ORGANIZATIONS:
        entry = connected.get(org["id"])
        seen.add(org["id"])
        listed.append({
            **org,
            "connected": bool(entry),
            "connection": _org_public_view(entry) if entry else None,
        })
    for org_id, entry in connected.items():
        if org_id in seen:
            continue
        listed.append({
            "id": org_id,
            "name": entry.get("org_name") or org_id,
            "hub_url": entry.get("hub_url"),
            "logo": None,
            "connected": True,
            "connection": _org_public_view(entry),
        })
    return {
        "organizations": listed,
        "public": {"id": PUBLIC_ORG_ID, "name": PUBLIC_ORG_NAME, "connected": True},
    }


@app.post("/api/org/connect")
async def org_connect(req: OrgConnectRequest, request: Request):
    """Start the hub handshake: mint a nonce, then open the hub's /connect in
    the user's browser pointed back at this backend's /org/callback."""
    _purge_expired_org_states()

    if req.hub_url:
        try:
            hub_url = normalize_hub_url(req.hub_url)
        except ValueError as e:
            raise HTTPException(status_code=400, detail=str(e))
        org_hint = org_hint_from_hub_url(hub_url)
    elif req.org_id:
        org = find_organization(req.org_id)
        if not org:
            raise HTTPException(status_code=404, detail=f"Unknown organization '{req.org_id}'")
        hub_url = org["hub_url"]
        org_hint = org["id"]
    else:
        raise HTTPException(status_code=400, detail="Supply an org_id or a hub_url")

    started = await _start_org_connect(hub_url, _callback_host(request), key=org_hint or hub_url)
    return {"status": "opened", "org_id": org_hint, **started}


@app.get("/org/callback")
async def org_callback(
    # Aliased because a parameter called `state` would shadow the module-level
    # AppState for the whole function body.
    state_nonce: str = Query("", alias="state"),
    app_id: str = "",
    app_key: str = "",
    gateway: str = "",
    org_id: str = "",
    org_name: str = "",
    email: str = "",
    expires: str = "",
    tables: str = "",
    error: str = "",
):
    """Receives the hub's redirect carrying a freshly minted personal
    credential, validates the nonce, and stores it 0600."""
    page = {
        "page_title": "VibeFoundry — Organization connected",
        "heading": "Connected!",
        "failed_heading": "Connection failed",
    }
    if error:
        return HTMLResponse(_auth_callback_html(error=error[:300], **page), status_code=400)

    _purge_expired_org_states()
    pending = _pending_org_states.pop(state_nonce, None) if state_nonce else None
    if not pending:
        return HTMLResponse(
            _auth_callback_html(error="Invalid or expired connection attempt. Try again from the IDE.", **page),
            status_code=400,
        )
    # The nonce came back, so this handshake is over whatever the hub sent with
    # it. Release the slot now or the next expiry inside the 10-minute window
    # would silently reuse a tab the user has already finished with.
    _clear_org_connect_inflight(state_nonce)
    if not (app_id and app_key and gateway and org_id):
        return HTMLResponse(
            _auth_callback_html(error="The hub's response was missing credential fields.", **page),
            status_code=400,
        )

    try:
        gateway_url = normalize_gateway_url(gateway)
    except ValueError:
        return HTMLResponse(_auth_callback_html(error="The hub sent an unusable gateway URL.", **page), status_code=400)

    try:
        table_count = int(tables)
    except (TypeError, ValueError):
        table_count = None

    orgs = _read_orgs()
    orgs[org_id] = {
        "org_id": org_id,
        "org_name": org_name or org_id,
        "hub_url": pending["hub_url"],
        "gateway": gateway_url,
        "app_id": app_id,
        "app_key": app_key,
        "email": email,
        "expires": expires,
        "connected_at": datetime.now(timezone.utc).isoformat(timespec="seconds").replace("+00:00", "Z"),
        "tables": table_count,
    }
    _write_orgs(orgs)
    # The user came back from connecting; show them what they just gained.
    state.open_org_catalog = True
    return HTMLResponse(_auth_callback_html(**page))


@app.get("/api/org/status")
async def org_status():
    """Which orgs are connected, as whom, and for how much longer. Never the key."""
    return {"organizations": [_org_public_view(e) for e in _read_orgs().values()]}


@app.post("/api/org/disconnect")
async def org_disconnect(req: OrgDisconnectRequest):
    return {"status": "ok", "org_id": req.org_id, "disconnected": _drop_org(req.org_id)}


@app.get("/api/org/catalog")
async def org_catalog(request: Request):
    """Every connected org's tables plus the public library, in one shape.

    One org being unreachable must not blank the catalogue — its failure is
    reported alongside the tables that did load, exactly as a missing public
    dataset doesn't blank the public page."""
    tables: list[dict] = []
    errors: list[dict] = []
    reauth: list[str] = []
    reauth_hub: Optional[str] = None

    for entry in _read_orgs().values():
        org_id = entry.get("org_id")
        try:
            res = await _gateway_request(entry, "GET", "/v1/tables")
            payload = _gateway_json(res)
        except OrgReauthRequired as e:
            reauth.append(org_id)
            reauth_hub = reauth_hub or e.hub_url
            continue
        except HTTPException as e:
            errors.append({"org_id": org_id, "error": str(e.detail)})
            continue
        for table in payload.get("tables", []):
            tables.append({
                "source": "org",
                "org_id": org_id,
                "org_name": entry.get("org_name"),
                "id": table.get("id"),
                "title": table.get("title") or table.get("id"),
                "rows": table.get("rows"),
                "columns": table.get("columns") or [],
            })

    try:
        manifest = await _public_manifest()
    except HTTPException as e:
        errors.append({"org_id": PUBLIC_ORG_ID, "error": str(e.detail)})
        manifest = {"datasets": []}

    entries = [d for d in manifest.get("datasets", []) if d.get("id")]
    async with httpx.AsyncClient(timeout=_DATA_TIMEOUT) as client:
        async def columns_for(did: str):
            # The manifest carries counts, not column names; the per-dataset
            # json is where those live. A missing one costs column names, not
            # the row.
            try:
                res = await client.get(f"{PUBLIC_DATA_BASE_URL}/{did}.json")
                res.raise_for_status()
                body = res.json()
                return [c.get("name") for c in body.get("columns", []) if c.get("name")]
            except Exception:
                return []

        column_lists = await asyncio.gather(*(columns_for(d["id"]) for d in entries))

    for dataset, columns in zip(entries, column_lists):
        tables.append({
            "source": "public",
            "org_id": PUBLIC_ORG_ID,
            "org_name": PUBLIC_ORG_NAME,
            "id": dataset.get("id"),
            "title": dataset.get("title") or dataset.get("id"),
            "rows": dataset.get("rowCount"),
            "columns": columns,
        })

    # `reauth_org_ids`, not `reauth_required`: the update below can set
    # top-level `status` to the string "reauth_required", and one key must not
    # mean both "the signal" and "the list of orgs it applies to".
    result = {"tables": tables, "errors": errors, "reauth_org_ids": reauth, "status": "ok"}
    if reauth:
        # The tables that did load are still returned alongside: the caller
        # retries the whole catalogue once the browser round trip lands, and a
        # frontend that only reads `tables` keeps working meanwhile.
        result.update(await _reauth_response(reauth[0], reauth_hub, _callback_host(request)))
    return result


@app.get("/api/org/schema/{org_id}/{table_id}")
async def org_schema(org_id: str, table_id: str, request: Request):
    """One table's column profile — what a caller reads before writing SQL."""
    if not _valid_table_id(table_id):
        raise HTTPException(status_code=400, detail="Invalid table_id")

    if org_id == PUBLIC_ORG_ID:
        async with httpx.AsyncClient(timeout=_DATA_TIMEOUT) as client:
            meta = await _public_dataset_meta(client, table_id)
        return {
            "source": "public",
            "org_id": PUBLIC_ORG_ID,
            "id": meta.get("id", table_id),
            "title": meta.get("title", table_id),
            "description": meta.get("description", ""),
            "rows": meta.get("rowCount"),
            "refreshedAt": meta.get("refreshedAt", ""),
            "columns": meta.get("columns", []),
        }

    try:
        entry = _require_org(org_id)
        res = await _gateway_request(entry, "GET", f"/v1/tables/{table_id}/schema")
        payload = _gateway_json(res)
    except OrgNotConnected as e:
        return _not_connected_response(e.org_id)
    except OrgReauthRequired as e:
        return await _reauth_response(e.org_id, e.hub_url, _callback_host(request))
    return {"source": "org", "org_id": org_id, **payload}


@app.post("/api/org/query")
async def org_query(req: OrgQueryRequest, request: Request):
    """Answer a question with SQL, at the gateway for an org and locally for
    public data. This is the path that means a question costs a few hundred
    rows rather than a whole-table download.

    A named `script_name` is resolved before any SQL runs - same policy as
    /api/org/pull: an unusable name is the caller's 400 at the door, never a
    result quietly redirected somewhere else after the work is done."""
    if req.script_name:
        _script_folders(req.script_name)
    if req.org_id == PUBLIC_ORG_ID:
        return await _run_public_sql(req.sql, req.limit, req.script_name)
    try:
        entry = _require_org(req.org_id)
        return await _run_org_sql(entry, req.sql, req.limit, req.script_name)
    except OrgNotConnected as e:
        return _not_connected_response(e.org_id)
    except OrgReauthRequired as e:
        return await _reauth_response(e.org_id, e.hub_url, _callback_host(request))


@app.post("/api/org/pull")
async def org_pull(req: OrgPullRequest, request: Request):
    """Land data on disk for a script to read — a query result when `sql` is
    given, the whole table when it isn't.

    With `script_name` the cut lands in that Track 0 script's raw_pulls/;
    without it, in input_folder/ for Tracks 1-4, exactly as before."""
    folder = _script_folders(req.script_name)["raw_pulls"] if req.script_name else _input_folder()

    if req.org_id == PUBLIC_ORG_ID:
        if req.sql:
            df, wanted, _ = await _collect_public_sql(req.sql, None)
            dest = _safe_dest(folder, req.filename or f"{_query_result_stem(wanted)}.parquet")
            df.write_parquet(dest)
        else:
            if not _valid_dataset_id(req.table_id):
                raise HTTPException(status_code=400, detail="Invalid table_id")
            # force=True: an explicit pull is the user asking for current data,
            # not for whatever copy is already sitting in input_folder.
            source, _ = await _ensure_public_parquet(req.table_id, force=True)
            dest = source
            # _ensure_public_parquet only ever writes to input_folder/, so a
            # Track 0 pull has to be copied on from there.
            if req.filename or folder.resolve() != source.parent.resolve():
                dest = _safe_dest(folder, req.filename or source.name)
                if dest != source:
                    shutil.copyfile(source, dest)
        await _after_download()
        return {
            "success": True,
            "org_id": PUBLIC_ORG_ID,
            "filename": dest.name,
            "bytes": dest.stat().st_size,
            "path": str(dest),
            "script_name": req.script_name,
        }

    if not req.sql and not _valid_table_id(req.table_id):
        raise HTTPException(status_code=400, detail="Invalid table_id")

    try:
        entry = _require_org(req.org_id)
        if req.sql:
            body = {"sql": req.sql, "format": "parquet"}
            res = await _gateway_request(entry, "POST", "/v1/query", json_body=body)
            default_name = f"{req.table_id or 'query'}_cut.parquet"
        else:
            res = await _gateway_request(entry, "GET", f"/v1/tables/{req.table_id}")
            default_name = f"{req.table_id}.parquet"
    except OrgNotConnected as e:
        return _not_connected_response(e.org_id)
    except OrgReauthRequired as e:
        return await _reauth_response(e.org_id, e.hub_url, _callback_host(request))

    _raise_for_gateway_error(res)

    dest = _safe_dest(folder, req.filename or default_name)
    dest.write_bytes(res.content)
    await _after_download()
    return {
        "success": True,
        "org_id": req.org_id,
        "table_id": req.table_id,
        "filename": dest.name,
        "bytes": dest.stat().st_size,
        "path": str(dest),
        "script_name": req.script_name,
    }


@app.get("/api/scripts")
async def list_scripts():
    """List available scripts"""
    if not state.project_folder:
        raise HTTPException(status_code=400, detail="No project folder selected")

    scripts_folder = state.project_folder / "app_folder" / "scripts"
    scripts = discover_scripts(scripts_folder)

    return {
        "scripts": [
            {
                "path": str(s),
                "relative_path": str(s.relative_to(scripts_folder)),
                "name": s.name
            }
            for s in scripts
        ]
    }


@app.post("/api/scripts/run")
async def run_scripts(request: RunScriptsRequest):
    """Run selected scripts"""
    import asyncio

    if not state.project_folder:
        raise HTTPException(status_code=400, detail="No project folder selected")

    results: list[ScriptResultResponse] = []

    for script_path in request.scripts:
        # Run in thread pool so server stays responsive (allows stop requests)
        result = await asyncio.to_thread(run_script, Path(script_path), state.project_folder)
        results.append(ScriptResultResponse(
            script_path=result.script_path,
            success=result.success,
            stdout=result.stdout,
            stderr=result.stderr,
            return_code=result.return_code,
            error=result.error,
            timed_out=result.timed_out,
            streamlit_url=result.streamlit_url
        ))

    # Regenerate metadata after running scripts (skip for .sh/.bat since they are long-running apps)
    ran_only_launchers = all(
        Path(s).suffix.lower() in (".sh", ".bat") for s in request.scripts
    )
    if not ran_only_launchers:
        generate_metadata(state.project_folder)

    return {"results": [r.model_dump() for r in results]}


class RunExternalRequest(BaseModel):
    scriptPath: str


@app.post("/api/scripts/run-external")
async def run_script_external(request: RunExternalRequest):
    """Launch a script in the system's external terminal."""
    import subprocess

    if not state.project_folder:
        raise HTTPException(status_code=400, detail="No project folder selected")

    script_path = Path(request.scriptPath)
    if not script_path.exists():
        raise HTTPException(status_code=404, detail="Script not found")

    cwd = str(state.project_folder)
    script = str(script_path)

    if sys.platform == 'darwin':
        # macOS: open Terminal.app and run the script
        apple_script = f'''tell application "Terminal"
    activate
    do script "cd {cwd} && python \\"{script}\\""
end tell'''
        subprocess.Popen(['osascript', '-e', apple_script])
    elif sys.platform == 'win32':
        # Windows: open cmd and run the script
        subprocess.Popen(
            f'start cmd /k "cd /d {cwd} && python \\"{script}\\""',
            shell=True
        )
    else:
        # Linux: try common terminal emulators
        for term_cmd in [
            ['gnome-terminal', '--', 'bash', '-c', f'cd {cwd} && python "{script}"; exec bash'],
            ['xterm', '-e', f'cd {cwd} && python "{script}"; bash'],
        ]:
            try:
                subprocess.Popen(term_cmd)
                break
            except FileNotFoundError:
                continue

    return {"success": True, "scriptPath": request.scriptPath}


@app.post("/api/scripts/stop")
async def stop_scripts():
    """Stop all currently running scripts"""
    stopped = stop_all_scripts()
    print(f"[Scripts] Stopped {stopped} running script(s)")
    return {"success": True, "stopped": stopped}


@app.get("/api/processes")
async def get_running_processes():
    """List all currently running script processes"""
    processes = list_running_processes()
    return {"processes": processes}


class StopProcessRequest(BaseModel):
    pid: int


@app.post("/api/processes/stop")
async def stop_single_process(request: StopProcessRequest):
    """Stop a specific process by PID"""
    success = stop_process(request.pid)
    if success:
        print(f"[Processes] Stopped process {request.pid}")
        return {"success": True, "pid": request.pid}
    else:
        return {"success": False, "error": f"Process {request.pid} not found or could not be stopped"}


@app.post("/api/metadata/generate")
async def regenerate_metadata():
    """Force metadata regeneration"""
    if not state.project_folder:
        raise HTTPException(status_code=400, detail="No project folder selected")

    input_meta, output_meta = await asyncio.to_thread(generate_metadata, state.project_folder)

    return {
        "success": True,
        "input_metadata": input_meta,
        "output_metadata": output_meta
    }


class PipInstallRequest(BaseModel):
    package: str


@app.post("/api/pip/install")
async def pip_install(request: PipInstallRequest):
    """Install a Python package using pip"""
    import subprocess
    import sys

    # Sanitize package name - only allow alphanumeric, hyphens, underscores, brackets
    package = request.package.strip()
    if not package or not all(c.isalnum() or c in '-_[],' for c in package):
        raise HTTPException(status_code=400, detail="Invalid package name")

    try:
        # Run pip install
        result = subprocess.run(
            [sys.executable, "-m", "pip", "install", package],
            capture_output=True,
            text=True,
            timeout=120  # 2 minute timeout
        )

        return {
            "success": result.returncode == 0,
            "package": package,
            "stdout": result.stdout,
            "stderr": result.stderr,
            "return_code": result.returncode
        }
    except subprocess.TimeoutExpired:
        return {
            "success": False,
            "package": package,
            "stdout": "",
            "stderr": "Installation timed out",
            "return_code": -1
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to install package: {str(e)}")


@app.get("/api/watch/check")
async def check_for_changes():
    """Manually check for file changes"""
    if not state.watcher:
        return {"changes": False}

    input_changes, output_changes, script_changes = state.watcher.check_once()

    has_changes = bool(input_changes or output_changes or script_changes)

    if input_changes or output_changes:
        await asyncio.to_thread(generate_metadata, state.project_folder)

    return {
        "changes": has_changes,
        "input_changes": [{"path": c.path, "type": c.change_type} for c in input_changes],
        "output_changes": [{"path": c.path, "type": c.change_type} for c in output_changes],
        "script_changes": [{"path": c.path, "type": c.change_type} for c in script_changes]
    }


# Filesystem browsing endpoints

@app.get("/api/fs/home")
async def get_home_directory():
    """Get user's home directory"""
    return {"path": str(Path.home())}


@app.get("/api/fs/list")
async def list_directory(path: str = ""):
    """List directories at a given path (for folder picker)"""
    if not path:
        path = str(Path.home())

    target = Path(path)

    if not target.exists():
        raise HTTPException(status_code=404, detail="Path does not exist")

    if not target.is_dir():
        raise HTTPException(status_code=400, detail="Path is not a directory")

    folders = []
    try:
        for item in sorted(target.iterdir()):
            # Only show directories, skip hidden files
            if item.is_dir() and not item.name.startswith('.'):
                folders.append({
                    "name": item.name,
                    "path": str(item)
                })
    except PermissionError:
        raise HTTPException(status_code=403, detail="Permission denied")

    return {
        "current": str(target),
        "parent": str(target.parent) if target.parent != target else None,
        "folders": folders
    }


class MkdirRequest(BaseModel):
    path: str
    name: str


@app.post("/api/fs/mkdir")
async def create_directory(request: MkdirRequest):
    """Create a new directory"""
    parent = Path(request.path)

    # If path is relative, make it relative to project folder
    if not parent.is_absolute() and state.project_folder:
        parent = state.project_folder / request.path

    if not parent.exists():
        raise HTTPException(status_code=404, detail=f"Parent path does not exist: {parent}")

    if not parent.is_dir():
        raise HTTPException(status_code=400, detail=f"Parent path is not a directory: {parent}")

    # Sanitize folder name - no path traversal
    name = request.name.strip()
    if not name or '/' in name or '\\' in name or name.startswith('.'):
        raise HTTPException(status_code=400, detail="Invalid folder name")

    new_folder = parent / name

    if new_folder.exists():
        raise HTTPException(status_code=409, detail=f"Folder already exists: {new_folder}")

    try:
        new_folder.mkdir(parents=False, exist_ok=False)
        return {"success": True, "path": str(new_folder)}
    except PermissionError:
        raise HTTPException(status_code=403, detail="Permission denied")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to create folder: {str(e)}")


# Directories to skip when building the file tree (heavy/irrelevant for the IDE)
TREE_BLACKLIST = {
    'node_modules', '__pycache__', '.next', 'build', 'dist',
    '.cache', '.parcel-cache', '.turbo', 'coverage',
    'env', 'venv', '.venv',
}


def _is_env_file(name: str) -> bool:
    """`.env` files are hidden dotfiles, but the IDE surfaces them anyway so
    users can paste their API keys straight into one. Matches `.env`,
    `.env.local`, `.env.production`, etc."""
    return name == '.env' or name.startswith('.env.')


def build_file_tree(path: Path, base_path: Path, deleted_files: list = None, in_app_folder: bool = False) -> dict:
    """Build a file tree recursively"""
    if deleted_files is None:
        deleted_files = []

    rel_path = path.relative_to(base_path).as_posix()
    is_file = path.is_file()
    node = {
        "name": path.name,
        "path": "" if rel_path == "." else rel_path,
        "isDirectory": not is_file,
        "extension": path.suffix if is_file else None,
        "lastModified": path.stat().st_mtime if is_file else None,
    }

    if path.is_dir():
        children = []
        # Check if we're entering app_folder
        entering_app_folder = in_app_folder or path.name == "app_folder"
        try:
            for item in sorted(path.iterdir()):
                # Skip hidden files — except .env files, where users paste keys
                if item.name.startswith('.') and not _is_env_file(item.name):
                    continue
                # Skip blacklisted directories
                if item.is_dir() and item.name in TREE_BLACKLIST:
                    continue

                children.append(build_file_tree(item, base_path, deleted_files, entering_app_folder))
        except PermissionError:
            pass
        node["children"] = children

    return node


@app.get("/api/files/tree")
async def get_file_tree():
    """Get the complete file tree for the project"""
    if not state.project_folder:
        raise HTTPException(status_code=400, detail="No project folder selected")

    deleted_files = []
    tree = await asyncio.to_thread(build_file_tree, state.project_folder, state.project_folder, deleted_files)
    return {"tree": tree, "deletedFiles": deleted_files}


@app.get("/api/files/read")
async def read_file(path: str, sheet: Optional[str] = None, asData: bool = False):
    """Read a file's content - streams from disk, doesn't hold data in memory"""
    if not state.project_folder:
        raise HTTPException(status_code=400, detail="No project folder selected")

    # Clear any previous DataFrame state
    if df_state.file_path is not None:
        df_state.clear()

    file_path = state.project_folder / path
    print(f"[File Read] Loading: {path}")

    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")

    if not file_path.is_file():
        raise HTTPException(status_code=400, detail="Path is not a file")

    # Security check - ensure path is within project folder
    try:
        file_path.resolve().relative_to(state.project_folder.resolve())
    except ValueError:
        raise HTTPException(status_code=403, detail="Access denied")

    # Determine file type and read accordingly
    ext = file_path.suffix.lower()
    binary_extensions = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.ico', '.webp', '.pdf', '.zip', '.tar', '.gz'}
    dataframe_extensions = {'.csv', '.xlsx', '.xls', '.parquet', '.geoparquet'}

    # Spreadsheets get shown as they look — styled cells and charts — rather
    # than flattened into a dataframe, which discards everything that makes a
    # spreadsheet a spreadsheet. asData=1 opts back into the grid for sorting
    # and scrolling large sheets.
    if ext in {'.xlsx', '.xlsm'} and not asData:
        return {
            "type": "spreadsheet",
            "filename": file_path.name,
            "path": str(file_path.relative_to(state.project_folder)),
        }

    if ext in dataframe_extensions:
        print(f"[File Read] Parsing dataframe: {path}")
        temp_file_path = None  # Track temp files for cleanup
        excel_sheet_names = None  # Track Excel sheet names
        excel_active_sheet = None

        # ── Try metadata first for column names & row count ──
        # This avoids scanning the file at all for basic info
        meta_columns = None
        meta_rows = None
        try:
            from vibefoundry.metadata import _metadata_cache
            cache_key = str(file_path)
            mtime = file_path.stat().st_mtime
            if cache_key in _metadata_cache:
                cached_mtime, cached_rows, cached_cols, cached_info = _metadata_cache[cache_key]
                if cached_mtime == mtime:
                    meta_columns = cached_cols
                    meta_rows = cached_rows
                    print(f"[File Read] Using cached metadata: {len(meta_columns)} cols, {meta_rows} rows")
        except Exception:
            pass

        # Also check the metadata txt files for row/column info
        if meta_columns is None:
            try:
                for meta_name in ("input_metadata.txt", "output_metadata.txt"):
                    meta_path = state.project_folder / "app_folder" / "meta_data" / meta_name
                    if meta_path.exists():
                        meta_text = meta_path.read_text(encoding="utf-8")
                        # Look for this file's entry in metadata
                        fname = file_path.name
                        if f"File: {fname}" in meta_text or str(file_path) in meta_text:
                            import re
                            # Extract row count
                            for line in meta_text.split('\n'):
                                if fname in line or str(file_path.relative_to(state.project_folder)) in line:
                                    # Found the file section, scan next lines
                                    pass
                            # Parse rows from metadata
                            sections = meta_text.split("File: ")
                            for section in sections:
                                if fname in section or str(file_path) in section:
                                    rows_match = re.search(r'Rows:\s*(\d+)', section)
                                    if rows_match:
                                        meta_rows = int(rows_match.group(1))
                                    cols_match = re.search(r'Columns\s*\((\d+)\):', section)
                                    if cols_match:
                                        # Extract column names from the "    - colname" lines
                                        col_lines = re.findall(r'^\s+-\s+(.+?)(?:\s+\[.*\])?$', section, re.MULTILINE)
                                        if col_lines:
                                            meta_columns = col_lines
                                    break
                            if meta_columns:
                                print(f"[File Read] Using metadata txt: {len(meta_columns)} cols, {meta_rows} rows")
                                break
            except Exception:
                pass

        try:
            if ext == '.csv':
                try:
                    # Read raw bytes to detect line endings and separator
                    with open(file_path, 'rb') as f:
                        sample = f.read(4096)

                    # Detect line ending style
                    has_crlf = b'\r\n' in sample
                    has_lf = b'\n' in sample
                    has_cr = b'\r' in sample

                    # Detect separator from first line
                    if has_crlf:
                        first_line = sample.split(b'\r\n')[0].decode('utf-8', errors='replace')
                    elif has_lf:
                        first_line = sample.split(b'\n')[0].decode('utf-8', errors='replace')
                    elif has_cr:
                        first_line = sample.split(b'\r')[0].decode('utf-8', errors='replace')
                    else:
                        first_line = sample.decode('utf-8', errors='replace')

                    # Detect separator
                    if '\t' in first_line:
                        separator = '\t'
                    elif ';' in first_line:
                        separator = ';'
                    else:
                        separator = ','

                    # Handle old Mac CR-only line endings - need temp file for streaming
                    needs_cr_conversion = has_cr and not has_lf and not has_crlf
                    actual_file_path = file_path

                    if needs_cr_conversion:
                        import tempfile
                        with open(file_path, 'rb') as f:
                            content = f.read()
                        content = content.replace(b'\r', b'\n')
                        tf = tempfile.NamedTemporaryFile(mode='wb', suffix='.csv', delete=False)
                        tf.write(content)
                        tf.close()
                        temp_file_path = tf.name
                        actual_file_path = Path(temp_file_path)
                        del content

                    # Store CSV file info for streaming
                    df_state.clear()
                    df_state.file_path = Path(actual_file_path).as_posix()
                    df_state.csv_separator = separator
                    df_state.file_type = 'csv'

                    # Get schema from a quick scan (only infer from first rows)
                    lf = pl.scan_csv(actual_file_path, separator=separator, infer_schema_length=10000)
                    df_state.columns = lf.collect_schema().names()
                    schema = lf.collect_schema()

                    # CSVs never go through massive file filtering — they get auto-converted
                    # to Parquet on upload if >50MB. Just load normally.
                    # Use metadata row count if available, otherwise count
                    if meta_rows is not None:
                        df_state.total_rows = meta_rows
                    else:
                        df_state.total_rows = lf.select(pl.len()).collect().item()

                except Exception as csv_err:
                    return {"type": "error", "message": f"Could not read CSV file: {csv_err}", "filename": file_path.name}

            elif ext in {'.parquet', '.geoparquet'}:
                try:
                    df_state.clear()
                    df_state.file_path = file_path.as_posix()
                    df_state.file_type = 'parquet'
                    df_state.csv_separator = ','

                    # Get row count cheaply from parquet metadata if not cached
                    parquet_rows = meta_rows
                    if parquet_rows is None:
                        try:
                            import pyarrow.parquet as pq
                            parquet_rows = pq.ParquetFile(file_path).metadata.num_rows
                        except Exception:
                            parquet_rows = 0

                    # For massive files, use pyarrow metadata to get columns/rows
                    # without scanning the data — avoids OOM on huge files
                    if is_file_massive(file_path, total_rows=parquet_rows):
                        try:
                            import pyarrow.parquet as pq
                            pf = pq.ParquetFile(file_path)
                            arrow_schema = pf.schema_arrow
                            df_state.columns = [f.name for f in arrow_schema]
                            df_state.total_rows = pf.metadata.num_rows
                            schema = {}
                            for field in arrow_schema:
                                schema[field.name] = str(field.type)
                            file_size = file_path.stat().st_size
                            profile_path = get_profile_cache_path(state.project_folder, file_path)
                            has_valid_profile = is_profile_valid(profile_path, file_path)
                            col_dtypes = {col: schema.get(col, "Unknown") for col in df_state.columns}
                            print(f"[File Read] MASSIVE file detected: {file_path.name} ({file_size / 1024 / 1024:.0f} MB). Profile valid: {has_valid_profile}")
                            return {
                                "type": "massive_file",
                                "filename": file_path.name,
                                "filePath": path,
                                "fileSize": file_size,
                                "columns": df_state.columns,
                                "totalRows": df_state.total_rows,
                                "hasProfile": has_valid_profile,
                                "columnDtypes": col_dtypes,
                            }
                        except Exception as pq_err:
                            return {"type": "error", "message": f"Could not read massive Parquet file metadata: {pq_err}", "filename": file_path.name}

                    try:
                        lf = pl.scan_parquet(file_path)
                        df_state.columns = lf.collect_schema().names()
                        schema = lf.collect_schema()
                        # Use metadata row count if available
                        if meta_rows is not None:
                            df_state.total_rows = meta_rows
                        else:
                            df_state.total_rows = lf.select(pl.len()).collect().item()
                    except Exception:
                        # Fallback to pyarrow (e.g. geoparquet with geometry columns)
                        try:
                            import pyarrow.parquet as pq

                            parquet_file = pq.ParquetFile(file_path)
                            arrow_schema = parquet_file.schema_arrow

                            valid_columns = []
                            for field in arrow_schema:
                                if hasattr(field.type, 'extension_name') and 'geo' in str(field.type.extension_name).lower():
                                    continue
                                valid_columns.append(field.name)

                            if not valid_columns:
                                return {"type": "error", "message": "GeoParquet file contains only geometry columns.", "filename": file_path.name}

                            table = parquet_file.read(columns=valid_columns)
                            temp_df = pl.from_arrow(table)

                            df_state.columns = temp_df.columns
                            schema = temp_df.schema
                            df_state.total_rows = meta_rows if meta_rows is not None else len(temp_df)
                            lf = temp_df.lazy()
                            del temp_df
                        except Exception as pyarrow_err:
                            return {"type": "error", "message": f"Could not read Parquet file: {pyarrow_err}", "filename": file_path.name}

                except Exception as parquet_err:
                    return {"type": "error", "message": f"Could not read Parquet file: {parquet_err}", "filename": file_path.name}

            else:
                # Excel (.xlsx, .xls)
                try:
                    from openpyxl import load_workbook
                    df_state.clear()
                    df_state.file_path = file_path.as_posix()
                    df_state.file_type = 'excel'
                    df_state.csv_separator = ','

                    # Get sheet names
                    wb = load_workbook(file_path, read_only=True)
                    excel_sheet_names = wb.sheetnames
                    wb.close()

                    # Read the requested sheet (or first sheet by default)
                    excel_active_sheet = sheet if sheet and sheet in excel_sheet_names else excel_sheet_names[0]
                    target_sheet = excel_active_sheet
                    temp_df = pl.read_excel(file_path, sheet_name=target_sheet)
                    df_state.columns = temp_df.columns
                    schema = temp_df.schema
                    df_state.total_rows = len(temp_df)
                    lf = temp_df.lazy()
                    del temp_df
                except Exception as excel_err:
                    return {"type": "error", "message": f"Could not read Excel file: {excel_err}. Make sure 'openpyxl' is installed (pip install openpyxl).", "filename": file_path.name}

            # Compute column info from the full dataset (Polars lazy scan — efficient)
            try:
                column_info = _compute_full_column_info(lf, df_state.columns, schema)
            except Exception:
                column_info = {col: {"type": "categorical", "values": [], "count": 0, "nullCount": 0, "blankCount": 0, "uniqueCount": 0} for col in df_state.columns}

            df_state.column_info = column_info

            # Get first chunk for initial preview
            CHUNK_SIZE = 200
            first_chunk, total_rows = df_state.get_rows(0, CHUNK_SIZE)

            print(f"[File Read] Fast preview: {df_state.total_rows} total rows, showing {len(first_chunk)}")

            result = {
                "type": "dataframe",
                "filePath": path,
                "columns": df_state.columns,
                "columnInfo": column_info,
                "data": first_chunk,
                "totalRows": df_state.total_rows,
                "offset": 0,
                "limit": CHUNK_SIZE,
                "filename": file_path.name
            }

            # Include sheet info for Excel files
            if excel_sheet_names:
                result["sheetNames"] = excel_sheet_names
                result["activeSheet"] = excel_active_sheet

            return result
        except Exception as e:
            return {"type": "error", "message": f"Unexpected error reading file: {e}", "filename": file_path.name}
        finally:
            # Clean up temp files
            if temp_file_path:
                try:
                    os.remove(temp_file_path)
                except OSError:
                    pass

    elif ext in binary_extensions:
        # Images - return metadata only, frontend uses /api/image endpoint for fast direct loading
        image_extensions = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.ico', '.webp'}
        if ext in image_extensions:
            return {"type": "image", "path": path, "filename": file_path.name, "extension": ext}
        # PDF files - metadata only, frontend uses /api/pdf endpoint for direct streaming
        if ext == '.pdf':
            return {"type": "pdf", "path": path, "filename": file_path.name}
        # Other binary files - still use base64
        import base64
        content = base64.b64encode(file_path.read_bytes()).decode('utf-8')
        return {"content": content, "encoding": "base64", "filename": file_path.name}
    elif ext == '.json':
        # JSON files - parse and return structured data
        try:
            import json
            content = file_path.read_text(encoding='utf-8')
            data = json.loads(content)
            return {"type": "json", "data": data, "filename": file_path.name}
        except (json.JSONDecodeError, UnicodeDecodeError) as e:
            return {"type": "error", "message": f"Failed to parse JSON: {str(e)}", "filename": file_path.name}
    elif ext == '.docx':
        # Word documents - parse with python-docx
        try:
            from docx import Document
            doc = Document(file_path)
            paragraphs = []
            for para in doc.paragraphs:
                style = para.style.name if para.style else ""
                text = para.text
                if text.strip():
                    paragraphs.append({"text": text, "style": style})

            # Also extract tables
            tables = []
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    rows.append([cell.text for cell in row.cells])
                if rows:
                    tables.append(rows)

            return {
                "type": "docx",
                "paragraphs": paragraphs,
                "tables": tables,
                "filename": file_path.name
            }
        except ImportError:
            return {"type": "error", "message": "Install python-docx to preview Word files: pip install python-docx", "filename": file_path.name}
        except Exception as e:
            return {"type": "error", "message": f"Could not read Word document: {e}", "filename": file_path.name}
    elif ext == '.doc':
        return {"type": "unknown", "message": "Legacy .doc format is not supported. Save as .docx to preview.", "filename": file_path.name}
    else:
        try:
            content = file_path.read_text(encoding='utf-8')
            return {"type": "text", "content": content, "encoding": "utf-8", "filename": file_path.name}
        except UnicodeDecodeError:
            import base64
            content = base64.b64encode(file_path.read_bytes()).decode('utf-8')
            return {"content": content, "encoding": "base64", "filename": file_path.name}


@app.get("/api/image")
async def get_image(path: str):
    """Serve image files directly as binary for fast loading"""
    if not state.project_folder:
        raise HTTPException(status_code=400, detail="No project folder selected")

    file_path = state.project_folder / path

    # Security check
    try:
        file_path.resolve().relative_to(state.project_folder.resolve())
    except ValueError:
        raise HTTPException(status_code=403, detail="Access denied")

    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Image not found")

    # Map extensions to media types
    ext = file_path.suffix.lower()
    media_types = {
        '.png': 'image/png',
        '.jpg': 'image/jpeg',
        '.jpeg': 'image/jpeg',
        '.gif': 'image/gif',
        '.bmp': 'image/bmp',
        '.ico': 'image/x-icon',
        '.webp': 'image/webp',
        '.svg': 'image/svg+xml',
    }

    media_type = media_types.get(ext, 'application/octet-stream')
    return FileResponse(file_path, media_type=media_type)


@app.get("/api/spreadsheet")
async def get_spreadsheet(path: str, sheet: Optional[str] = None):
    """A spreadsheet reconstructed as it looks: styled cells, plus its charts.

    Returns {html, sheets, activeSheet, charts}. The html carries real fills,
    fonts, borders and merges; each chart is a definition read out of the
    workbook for the client to draw. Nothing here shells out — this is the path
    that works on a machine where no office suite can be installed.
    """
    if not state.project_folder:
        raise HTTPException(status_code=400, detail="No project folder selected")

    file_path = state.project_folder / path
    try:
        file_path.resolve().relative_to(state.project_folder.resolve())
    except ValueError:
        raise HTTPException(status_code=403, detail="Access denied")
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")

    result = await asyncio.to_thread(xlsx_view.render, file_path, sheet)
    if result is None:
        raise HTTPException(status_code=500, detail="Could not read this spreadsheet.")
    return result


@app.get("/api/pdf")
async def get_pdf(path: str):
    """Serve PDF files directly with application/pdf media type for inline iframe rendering."""
    if not state.project_folder:
        raise HTTPException(status_code=400, detail="No project folder selected")

    file_path = state.project_folder / path

    try:
        file_path.resolve().relative_to(state.project_folder.resolve())
    except ValueError:
        raise HTTPException(status_code=403, detail="Access denied")

    if not file_path.exists() or file_path.suffix.lower() != ".pdf":
        raise HTTPException(status_code=404, detail="PDF not found")

    return FileResponse(
        file_path,
        media_type="application/pdf",
        headers={"Content-Disposition": f'inline; filename="{file_path.name}"'},
    )


class WriteFileRequest(BaseModel):
    path: str
    content: str


@app.post("/api/files/write")
async def write_file(request: WriteFileRequest):
    """Write content to a file"""
    if not state.project_folder:
        raise HTTPException(status_code=400, detail="No project folder selected")

    file_path = state.project_folder / request.path

    # Security check - ensure path is within project folder
    try:
        file_path.resolve().relative_to(state.project_folder.resolve())
    except ValueError:
        raise HTTPException(status_code=403, detail="Access denied")

    # Create parent directories if needed
    file_path.parent.mkdir(parents=True, exist_ok=True)

    file_path.write_text(request.content, encoding='utf-8')

    return {"success": True, "path": request.path}


UPLOAD_CHUNK_SIZE = 8 * 1024 * 1024  # 8MB


SKIP_UPLOAD_NAMES = {".DS_Store", "Thumbs.db", "desktop.ini", ".localized"}


@app.post("/api/files/upload")
async def upload_file(
    file: UploadFile = File(...),
    folder: str = Form(...),
    relativePath: str = Form(default=""),
):
    """Upload a binary file to a folder, streaming to disk in chunks.

    `folder` is the project-relative destination directory (the folder the user
    right-clicked on). `relativePath` is the file's path relative to that
    directory — for a "Add Folder" upload it's the picker's webkitRelativePath
    (e.g. "MyData/sub/file.txt"), so the directory structure is preserved
    server-side. Older callers that don't send relativePath fall back to
    file.filename and land the file directly inside `folder`.
    """
    if not state.project_folder:
        raise HTTPException(status_code=400, detail="No project folder selected")

    if file.filename in SKIP_UPLOAD_NAMES:
        return {"success": True, "path": None, "skipped": True}

    base = folder.strip("/")
    rel = (relativePath or file.filename or "").lstrip("/")

    if not rel:
        raise HTTPException(status_code=400, detail="Missing file path")

    # Reject any path traversal in either piece before joining.
    for piece in (base, rel):
        if ".." in Path(piece).parts:
            raise HTTPException(status_code=400, detail="Invalid path")

    target_path = state.project_folder / base / rel if base else state.project_folder / rel

    # Belt-and-suspenders: even after the explicit `..` check, confirm the
    # resolved path is still inside the project folder.
    try:
        target_path.resolve().relative_to(state.project_folder.resolve())
    except ValueError:
        raise HTTPException(status_code=403, detail="Access denied")

    target_path.parent.mkdir(parents=True, exist_ok=True)

    # Stream file to disk in chunks to avoid loading entire file into memory
    with open(target_path, "wb") as f:
        while True:
            chunk = await file.read(UPLOAD_CHUNK_SIZE)
            if not chunk:
                break
            f.write(chunk)

    result_path = str(target_path.relative_to(state.project_folder))
    return {"success": True, "path": result_path, "converted": False}


class ConvertToParquetRequest(BaseModel):
    path: str
    deleteOriginal: bool = True


@app.post("/api/files/convert-to-parquet")
async def convert_to_parquet(request: ConvertToParquetRequest):
    """Convert a CSV or Excel file to Parquet alongside it. User-triggered only."""
    if not state.project_folder:
        raise HTTPException(status_code=400, detail="No project folder selected")

    source_path = state.project_folder / request.path
    try:
        source_path.resolve().relative_to(state.project_folder.resolve())
    except ValueError:
        raise HTTPException(status_code=403, detail="Access denied")

    if not source_path.exists() or not source_path.is_file():
        raise HTTPException(status_code=404, detail="File not found")

    ext = source_path.suffix.lower()
    if ext not in {".csv", ".xlsx", ".xls"}:
        raise HTTPException(status_code=400, detail="Only .csv, .xlsx, and .xls can be converted")

    parquet_path = source_path.with_suffix(".parquet")
    if parquet_path.exists():
        raise HTTPException(status_code=409, detail=f"{parquet_path.name} already exists")

    try:
        if ext == ".csv":
            print(f"[Convert] CSV → Parquet: {source_path.name}")
            pl.scan_csv(
                str(source_path),
                infer_schema_length=10000,
                null_values=["null", "NULL", "None", ""],
            ).sink_parquet(str(parquet_path))
        else:
            print(f"[Convert] Excel → Parquet: {source_path.name}")
            from openpyxl import load_workbook
            wb = load_workbook(source_path, read_only=True)
            sheet_name = wb.sheetnames[0]
            wb.close()
            pl.read_excel(source_path, sheet_name=sheet_name).write_parquet(str(parquet_path))
    except Exception as e:
        if parquet_path.exists():
            parquet_path.unlink()
        raise HTTPException(status_code=500, detail=f"Conversion failed: {e}")

    if request.deleteOriginal:
        try:
            source_path.unlink()
        except Exception as e:
            print(f"[Convert] Could not delete original after convert: {e}")

    rel_parent = Path(request.path).parent.as_posix()
    result_path = f"{rel_parent}/{parquet_path.name}" if rel_parent not in ("", ".") else parquet_path.name
    return {"success": True, "path": result_path}


class DeleteFileRequest(BaseModel):
    path: str
    isDirectory: bool = False


@app.post("/api/files/delete")
async def delete_file(request: DeleteFileRequest):
    """Delete a file or directory"""
    if not state.project_folder:
        raise HTTPException(status_code=400, detail="No project folder selected")

    file_path = state.project_folder / request.path

    # Security check - ensure path is within project folder
    try:
        file_path.resolve().relative_to(state.project_folder.resolve())
    except ValueError:
        raise HTTPException(status_code=403, detail="Access denied")

    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")

    import shutil

    # Collect files to clean up profiles for (before deletion)
    files_to_clean = []
    try:
        if request.isDirectory:
            for root, _dirs, fnames in os.walk(file_path):
                for fname in fnames:
                    files_to_clean.append(Path(os.path.join(root, fname)))
            shutil.rmtree(file_path)
        else:
            files_to_clean.append(file_path)
            file_path.unlink()
    except PermissionError:
        raise HTTPException(status_code=409, detail="Your File Is Still Open! Close It Before Deleting")

    # Remove any cached profile files for the deleted files
    for f in files_to_clean:
        profile_path = get_profile_cache_path(state.project_folder, f)
        if profile_path.exists():
            profile_path.unlink()
        meta_json = profile_path.with_suffix(".meta.json")
        if meta_json.exists():
            meta_json.unlink()

    # Regenerate metadata so profile files reflect the deletion
    generate_metadata(state.project_folder)

    return {"success": True, "path": request.path}


class RevealFileRequest(BaseModel):
    path: str


@app.post("/api/files/reveal")
async def reveal_file(request: RevealFileRequest):
    """Reveal a file or folder in the OS file manager (desktop only).

    macOS/Windows highlight the item itself; Linux has no portable
    "select this file" command, so it opens the containing folder.
    """
    import subprocess

    if not state.project_folder:
        raise HTTPException(status_code=400, detail="No project folder selected")

    file_path = state.project_folder / request.path

    # Security check - ensure path is within project folder
    try:
        resolved = file_path.resolve()
        resolved.relative_to(state.project_folder.resolve())
    except ValueError:
        raise HTTPException(status_code=403, detail="Access denied")

    if not resolved.exists():
        raise HTTPException(status_code=404, detail="File not found")

    try:
        if sys.platform == 'darwin':
            # -R reveals and selects the item in Finder
            subprocess.run(["open", "-R", str(resolved)], check=True)
        elif sys.platform == 'win32':
            # /select, opens Explorer with the item highlighted.
            # explorer.exe returns exit code 1 even on success, so no check.
            subprocess.run(["explorer", f"/select,{resolved}"])
        else:
            # Linux: no portable per-file select — open the containing folder
            target = resolved if resolved.is_dir() else resolved.parent
            subprocess.run(["xdg-open", str(target)], check=True)
    except (subprocess.CalledProcessError, FileNotFoundError) as e:
        raise HTTPException(status_code=500, detail=f"Could not open file manager: {e}")

    return {"success": True, "path": request.path}


class RenameRequest(BaseModel):
    oldPath: str
    newName: str


@app.post("/api/files/rename")
async def rename_file(request: RenameRequest):
    """Rename a file or directory"""
    if not state.project_folder:
        raise HTTPException(status_code=400, detail="No project folder selected")

    old_path = Path(request.oldPath)
    if not old_path.is_absolute():
        old_path = state.project_folder / request.oldPath

    # Security check
    try:
        old_path.resolve().relative_to(state.project_folder.resolve())
    except ValueError:
        raise HTTPException(status_code=403, detail="Access denied")

    if not old_path.exists():
        raise HTTPException(status_code=404, detail="File not found")

    new_path = old_path.parent / request.newName

    # Check if new path already exists
    if new_path.exists():
        raise HTTPException(status_code=400, detail="A file with that name already exists")

    import shutil
    shutil.move(str(old_path), str(new_path))

    # Regenerate metadata so profile files reflect the rename
    generate_metadata(state.project_folder)

    return {"success": True, "oldPath": str(old_path), "newPath": str(new_path)}


class MoveRequest(BaseModel):
    sourcePath: str
    destPath: str


@app.post("/api/files/move")
async def move_file(request: MoveRequest):
    """Move a file or directory"""
    if not state.project_folder:
        raise HTTPException(status_code=400, detail="No project folder selected")

    source_path = Path(request.sourcePath)
    dest_path = Path(request.destPath)

    if not source_path.is_absolute():
        source_path = state.project_folder / request.sourcePath
    if not dest_path.is_absolute():
        dest_path = state.project_folder / request.destPath

    # Security check
    try:
        source_path.resolve().relative_to(state.project_folder.resolve())
        dest_path.resolve().relative_to(state.project_folder.resolve())
    except ValueError:
        raise HTTPException(status_code=403, detail="Access denied")

    if not source_path.exists():
        raise HTTPException(status_code=404, detail="Source file not found")

    # Ensure destination directory exists
    dest_path.parent.mkdir(parents=True, exist_ok=True)

    import shutil
    shutil.move(str(source_path), str(dest_path))

    # Regenerate metadata so profile files reflect the move
    generate_metadata(state.project_folder)

    return {"success": True, "sourcePath": str(source_path), "destPath": str(dest_path)}


# DataFrame streaming endpoints

class DataFrameQueryRequest(BaseModel):
    filePath: str
    filters: dict = {}
    sort: Optional[dict] = None  # {column: str, direction: "asc"|"desc"}


@app.get("/api/dataframe/rows")
async def get_dataframe_rows(
    filePath: str,
    offset: int = 0,
    limit: int = 200
):
    """Get paginated rows - streams from disk, doesn't hold full file in memory"""
    if df_state.file_path is None:
        raise HTTPException(status_code=400, detail="No DataFrame loaded. Read a file first.")

    # Stream rows from disk
    rows, total_rows = df_state.get_rows(offset, limit)

    return {
        "data": rows,
        "offset": offset,
        "limit": limit,
        "totalRows": total_rows
    }


@app.post("/api/dataframe/query")
async def query_dataframe(request: DataFrameQueryRequest):
    """Apply filters and/or sort to the DataFrame - streams from disk"""
    if df_state.file_path is None:
        raise HTTPException(status_code=400, detail="DataFrame not loaded. Read the file first.")

    # Check if the requested file matches the loaded file (compare by filename since paths may differ)
    loaded_filename = Path(df_state.file_path).name
    requested_filename = Path(request.filePath).name
    if loaded_filename != requested_filename:
        raise HTTPException(status_code=400, detail=f"Different file loaded. Expected {requested_filename}, got {loaded_filename}")

    # Update filters and sort on state
    df_state.current_filters = request.filters
    df_state.current_sort = request.sort
    df_state.invalidate_filter_cache()  # Force recount

    # Get first chunk using streaming
    CHUNK_SIZE = 200
    rows, total_rows = df_state.get_rows(0, CHUNK_SIZE)

    # Compute cascading columnInfo from filtered data
    # For efficiency, we sample a limited number of rows for column stats
    cascading_column_info = await _compute_cascading_column_info()

    return {
        "data": rows,
        "totalRows": total_rows,
        "offset": 0,
        "limit": CHUNK_SIZE,
        "appliedFilters": request.filters,
        "appliedSort": request.sort,
        "columnInfo": cascading_column_info
    }


async def _compute_cascading_column_info() -> dict:
    """Compute column info from filtered data using optimized batched queries."""
    if df_state.file_path is None:
        return {}

    lf = df_state._get_lazy_frame()
    if lf is None:
        return {}

    lf = df_state._apply_filters_sort(lf)
    schema = lf.collect_schema()

    return _compute_column_info(lf, df_state.columns, schema)


@app.post("/api/dataframe/clear")
async def clear_dataframe():
    """Clear the DataFrame from memory"""
    df_state.clear()
    return {"success": True}


# --- Large file profiling endpoints ---

class ProfileRequest(BaseModel):
    filePath: str

class EstimateRequest(BaseModel):
    filePath: str
    filters: dict = {}

class FilteredPreviewRequest(BaseModel):
    filePath: str
    filters: dict = {}
    rowLimit: Optional[int] = None


@app.post("/api/dataframe/profile")
async def start_profile(request: ProfileRequest):
    """Start profiling a massive file. Progress is sent via WebSocket."""
    global _profiling_task

    if not state.project_folder:
        raise HTTPException(status_code=400, detail="No project folder selected")

    file_path = state.project_folder / request.filePath
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")

    ext = file_path.suffix.lower()
    if ext == ".csv":
        file_type = "csv"
    elif ext in {".parquet", ".geoparquet"}:
        file_type = "parquet"
    else:
        raise HTTPException(status_code=400, detail="Unsupported file type for profiling")

    # Check if profile already exists and is valid
    profile_path = get_profile_cache_path(state.project_folder, file_path)
    if is_profile_valid(profile_path, file_path):
        profile = read_cached_profile(profile_path)
        return {"status": "complete", "profile": profile}

    # Cancel any existing profiling task
    if _profiling_task and not _profiling_task.done():
        _profiling_task.cancel()

    async def _run_profiling():
        """Run profiling in a thread and push progress via WebSocket."""
        last_progress = [0]

        def on_progress(done: int, total: int):
            last_progress[0] = done
            # We'll send progress from the async wrapper below
            pass

        # Run the CPU-heavy profiling in a thread
        result = await asyncio.to_thread(
            profile_large_file, file_path, file_type, state.project_folder, on_progress
        )

        # Send completion via WebSocket
        msg = json.dumps({"type": "profile_complete", "filePath": request.filePath, "profile": result})
        disconnected = []
        for client in state.websocket_clients:
            try:
                await client.send_text(msg)
            except Exception:
                disconnected.append(client)
        for client in disconnected:
            state.websocket_clients.remove(client)

    # Start profiling with progress polling
    async def _profile_with_progress():
        """Wrapper that runs profiling and sends periodic progress updates."""
        progress_state = {"done": 0, "total": 1}

        def on_progress(done: int, total: int):
            progress_state["done"] = done
            progress_state["total"] = total

        # Start the profiling in a thread
        import concurrent.futures
        loop = asyncio.get_event_loop()
        with concurrent.futures.ThreadPoolExecutor(max_workers=1) as pool:
            future = loop.run_in_executor(
                pool, profile_large_file, file_path, file_type,
                state.project_folder, on_progress,
            )

            # Poll and send progress while profiling runs
            while not future.done():
                await asyncio.sleep(0.5)
                msg = json.dumps({
                    "type": "profile_progress",
                    "filePath": request.filePath,
                    "done": progress_state["done"],
                    "total": progress_state["total"],
                })
                disconnected = []
                for client in state.websocket_clients:
                    try:
                        await client.send_text(msg)
                    except Exception:
                        disconnected.append(client)
                for client in disconnected:
                    state.websocket_clients.remove(client)

            # Get result (may raise if profiling failed)
            result = await future

        # Send completion
        msg = json.dumps({
            "type": "profile_complete",
            "filePath": request.filePath,
            "profile": result,
        })
        disconnected = []
        for client in state.websocket_clients:
            try:
                await client.send_text(msg)
            except Exception:
                disconnected.append(client)
        for client in disconnected:
            state.websocket_clients.remove(client)

    _profiling_task = asyncio.create_task(_profile_with_progress())

    return {"status": "profiling", "message": "Profiling started. Progress will be sent via WebSocket."}


@app.get("/api/dataframe/profile/result")
async def get_profile_result(filePath: str):
    """Get the cached profile for a file."""
    if not state.project_folder:
        raise HTTPException(status_code=400, detail="No project folder selected")

    file_path = state.project_folder / filePath
    profile_path = get_profile_cache_path(state.project_folder, file_path)

    if not is_profile_valid(profile_path, file_path):
        raise HTTPException(status_code=404, detail="No valid profile found. Run profiling first.")

    profile = read_cached_profile(profile_path)
    return {"profile": profile}


@app.post("/api/dataframe/estimate-rows")
async def estimate_rows(request: EstimateRequest):
    """Estimate row count after applying filters. Fast on Parquet via predicate pushdown."""
    if not state.project_folder:
        raise HTTPException(status_code=400, detail="No project folder selected")

    file_path = state.project_folder / request.filePath
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")

    ext = file_path.suffix.lower()
    file_type = "csv" if ext == ".csv" else "parquet"
    separator = _detect_csv_separator(file_path) if file_type == "csv" else ","

    try:
        count = await asyncio.to_thread(
            estimate_filtered_rows, file_path, file_type, request.filters, separator
        )
        return {"estimatedRows": count}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Estimation failed: {e}")


@app.post("/api/dataframe/filtered-preview")
async def filtered_preview(request: FilteredPreviewRequest):
    """Load a filtered subset of a massive file for preview.
    Sets up df_state so subsequent /api/dataframe/rows calls work normally."""
    if not state.project_folder:
        raise HTTPException(status_code=400, detail="No project folder selected")

    file_path = state.project_folder / request.filePath
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")

    ext = file_path.suffix.lower()

    # Set up df_state
    df_state.clear()
    df_state.file_path = file_path.as_posix()
    if ext == ".csv":
        df_state.file_type = "csv"
        df_state.csv_separator = _detect_csv_separator(file_path)
    elif ext in {".parquet", ".geoparquet"}:
        df_state.file_type = "parquet"
    else:
        raise HTTPException(status_code=400, detail="Unsupported file type")

    lf = df_state._get_lazy_frame()
    if lf is None:
        raise HTTPException(status_code=500, detail="Failed to open file")

    df_state.columns = lf.collect_schema().names()
    schema = lf.collect_schema()

    # Apply the user's filters
    df_state.current_filters = request.filters
    df_state.invalidate_filter_cache()

    # Get total filtered rows
    # Apply row limit if specified
    df_state.row_limit = request.rowLimit if request.rowLimit else None
    df_state.invalidate_filter_cache()

    filtered_lf = df_state._apply_filters_sort(lf)
    if df_state.row_limit:
        filtered_lf = filtered_lf.slice(0, df_state.row_limit)
    total_filtered = filtered_lf.select(pl.len()).collect().item()
    df_state.total_rows = total_filtered

    # Compute column info on the filtered lazy frame
    try:
        column_info = _compute_column_info(filtered_lf, df_state.columns, schema)
    except Exception:
        column_info = {col: {"type": "categorical", "values": [], "count": 0, "nullCount": 0, "blankCount": 0, "uniqueCount": 0} for col in df_state.columns}
    df_state.column_info = column_info

    # Get first chunk
    CHUNK_SIZE = 200
    rows, total_rows = df_state.get_rows(0, CHUNK_SIZE)

    print(f"[Filtered Preview] {file_path.name}: {total_filtered} rows after filters")

    return {
        "type": "dataframe",
        "filePath": request.filePath,
        "columns": df_state.columns,
        "columnInfo": column_info,
        "data": rows,
        "totalRows": total_rows,
        "offset": 0,
        "limit": CHUNK_SIZE,
        "filename": file_path.name,
    }


# WebSocket for real-time updates

@app.websocket("/ws/watch")
async def websocket_watch(websocket: WebSocket):
    """WebSocket for file change notifications"""
    await websocket.accept()
    state.websocket_clients.append(websocket)

    try:
        while True:
            # Keep connection alive, wait for messages
            try:
                data = await asyncio.wait_for(websocket.receive_text(), timeout=30)
                # Handle any incoming messages (e.g., ping)
                if data == "ping":
                    await websocket.send_text("pong")
            except asyncio.TimeoutError:
                # Send keepalive
                await websocket.send_text('{"type": "keepalive"}')
    except WebSocketDisconnect:
        state.websocket_clients.remove(websocket)
    except Exception:
        if websocket in state.websocket_clients:
            state.websocket_clients.remove(websocket)


async def notify_data_change():
    """Notify all WebSocket clients of data change"""
    if state.project_folder:
        await asyncio.to_thread(generate_metadata, state.project_folder)

    message = '{"type": "data_change"}'
    disconnected = []

    for client in state.websocket_clients:
        try:
            await client.send_text(message)
        except Exception:
            disconnected.append(client)

    for client in disconnected:
        state.websocket_clients.remove(client)


async def notify_script_change(script_path: Path):
    """Notify all WebSocket clients of script change"""
    # Send full absolute path (same format as /api/scripts endpoint)
    full_path = str(script_path)
    # Use forward slashes for consistency on Windows
    full_path = full_path.replace("\\", "/")

    # Debounce: skip if we notified about this script in the last 3 seconds
    # Use lowercase key for case-insensitive matching (Windows paths)
    debounce_key = full_path.lower()
    now = time.time()
    if debounce_key in state.last_script_change:
        if now - state.last_script_change[debounce_key] < 3.0:
            print(f"[Script Change] Debounced (duplicate within 3s): {full_path}")
            return
    state.last_script_change[debounce_key] = now
    # Clean up old entries
    state.last_script_change = {k: v for k, v in state.last_script_change.items() if now - v < 10.0}

    print(f"[Script Change] Notifying {len(state.websocket_clients)} clients: {full_path}")
    message = json.dumps({"type": "script_change", "path": full_path})
    disconnected = []

    for client in state.websocket_clients:
        try:
            await client.send_text(message)
        except Exception:
            disconnected.append(client)

    for client in disconnected:
        state.websocket_clients.remove(client)


async def notify_output_file_change(file_path: Path, change_type: str):
    """Notify all WebSocket clients of output file change for auto-preview"""
    # Get relative path from project folder
    rel_path = str(file_path)
    if state.project_folder:
        try:
            rel_path = str(file_path.relative_to(state.project_folder))
        except ValueError:
            pass
    # Use forward slashes for consistency (Windows fix)
    rel_path = rel_path.replace("\\", "/")

    print(f"[Output Change] Notifying {len(state.websocket_clients)} clients: {rel_path}")
    message = json.dumps({"type": "output_file_change", "path": rel_path, "change_type": change_type})
    disconnected = []

    for client in state.websocket_clients:
        try:
            await client.send_text(message)
        except Exception:
            disconnected.append(client)

    for client in disconnected:
        state.websocket_clients.remove(client)


# Local Terminal WebSocket

def set_terminal_size(fd, rows, cols):
    """Set terminal window size"""
    if sys.platform == 'win32':
        return  # Not supported on Windows
    winsize = struct.pack("HHHH", rows, cols, 0, 0)
    fcntl.ioctl(fd, termios.TIOCSWINSZ, winsize)


@app.websocket("/ws/terminal")
async def websocket_terminal(websocket: WebSocket):
    """WebSocket for local terminal"""
    await websocket.accept()

    # Terminal not supported on Windows
    if sys.platform == 'win32':
        await websocket.send_text("Terminal not supported on Windows.\r\n")
        await websocket.close()
        return

    # Fork a PTY
    pid, fd = pty.fork()

    if pid == 0:
        # Child process - create new session/process group so we can kill all children
        os.setsid()
        cwd = str(state.project_folder) if state.project_folder else str(Path.home())
        os.chdir(cwd)
        os.environ["TERM"] = "xterm-256color"
        os.execvp("bash", ["bash", "-l"])
    else:
        # Parent process - relay data
        print(f"[Terminal] Started PTY process {pid}")
        set_terminal_size(fd, 24, 80)

        # Make fd non-blocking
        flags = fcntl.fcntl(fd, fcntl.F_GETFL)
        fcntl.fcntl(fd, fcntl.F_SETFL, flags | os.O_NONBLOCK)

        try:
            while True:
                # Check for data from terminal (non-blocking)
                r, _, _ = select.select([fd], [], [], 0.05)
                if fd in r:
                    try:
                        data = os.read(fd, 8192)
                        if data:
                            await websocket.send_text(data.decode("utf-8", errors="replace"))
                    except OSError:
                        break

                # Check for data from websocket (with timeout)
                try:
                    data = await asyncio.wait_for(websocket.receive_text(), timeout=0.05)
                    if data:
                        # Check for JSON commands
                        if data.startswith('{'):
                            try:
                                msg = json.loads(data)
                                if msg.get('type') == 'resize':
                                    rows = msg.get('rows', 24)
                                    cols = msg.get('cols', 80)
                                    set_terminal_size(fd, rows, cols)
                                elif msg.get('type') == 'ping':
                                    await websocket.send_text('{"type":"pong"}')
                            except json.JSONDecodeError:
                                pass
                        else:
                            os.write(fd, data.encode("utf-8"))
                except asyncio.TimeoutError:
                    pass
                except WebSocketDisconnect:
                    print(f"[Terminal] WebSocket disconnected, cleaning up PTY {pid}")
                    break
        finally:
            # Clean up: close fd and kill the entire process group
            print(f"[Terminal] Cleaning up PTY process {pid}")
            try:
                os.close(fd)
            except OSError:
                pass

            # Kill the entire process group (bash + all child processes like claude)
            try:
                # First try SIGTERM to the process group
                os.killpg(pid, signal.SIGTERM)
            except OSError:
                # Process group might not exist, try killing just the pid
                try:
                    os.kill(pid, signal.SIGTERM)
                except OSError:
                    pass

            # Give processes a moment to terminate gracefully
            await asyncio.sleep(0.5)

            # Force kill if still running
            try:
                os.killpg(pid, signal.SIGKILL)
            except OSError:
                try:
                    os.kill(pid, signal.SIGKILL)
                except OSError:
                    pass

            # Reap zombie process
            try:
                os.waitpid(pid, os.WNOHANG)
            except OSError:
                pass

            print(f"[Terminal] PTY process {pid} cleaned up")


# Serve static files (React app)

@app.get("/")
async def serve_index():
    """Serve the React app index.html"""
    static_dir = get_static_dir()
    index_path = static_dir / "index.html"

    if not index_path.exists():
        return JSONResponse(
            status_code=503,
            content={
                "error": "Frontend not built",
                "message": "Run 'npm run build' in the frontend directory first"
            }
        )

    return FileResponse(index_path)


# Serve root-level static files (icon.svg, manifest.json, sw.js, etc.)
@app.get("/icon.svg")
async def serve_icon():
    """Serve the favicon SVG"""
    static_dir = get_static_dir()
    icon_path = static_dir / "icon.svg"
    if icon_path.exists():
        return FileResponse(icon_path, media_type="image/svg+xml")
    return JSONResponse(status_code=404, content={"error": "icon.svg not found"})


@app.get("/icon-192.png")
async def serve_icon_192():
    static_dir = get_static_dir()
    path = static_dir / "icon-192.png"
    if path.exists():
        return FileResponse(path, media_type="image/png")
    return JSONResponse(status_code=404, content={"error": "icon-192.png not found"})


@app.get("/icon-512.png")
async def serve_icon_512():
    static_dir = get_static_dir()
    path = static_dir / "icon-512.png"
    if path.exists():
        return FileResponse(path, media_type="image/png")
    return JSONResponse(status_code=404, content={"error": "icon-512.png not found"})


@app.get("/manifest.json")
async def serve_manifest():
    static_dir = get_static_dir()
    path = static_dir / "manifest.json"
    if path.exists():
        return FileResponse(path, media_type="application/manifest+json")
    return JSONResponse(status_code=404, content={"error": "manifest.json not found"})


@app.get("/sw.js")
async def serve_service_worker():
    static_dir = get_static_dir()
    path = static_dir / "sw.js"
    if path.exists():
        return FileResponse(path, media_type="application/javascript")
    return JSONResponse(status_code=404, content={"error": "sw.js not found"})


@app.get("/vf_logo.png")
async def serve_vf_logo():
    static_dir = get_static_dir()
    path = static_dir / "vf_logo.png"
    if path.exists():
        return FileResponse(path, media_type="image/png")
    return JSONResponse(status_code=404, content={"error": "vf_logo.png not found"})


# Mount static files for assets (at module load time)
_static_dir = get_static_dir()
_assets_dir = _static_dir / "assets"
if _assets_dir.exists():
    app.mount("/assets", StaticFiles(directory=_assets_dir), name="assets")


def create_app() -> FastAPI:
    """Factory function for creating the app"""
    return app
